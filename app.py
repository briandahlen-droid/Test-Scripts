"""
Development Services Proposal Generator
Streamlit web application for generating professional proposal documents
Integrated with comprehensive property lookup for Pinellas, Hillsborough, and Pasco counties
"""
import streamlit as st
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import requests
from bs4 import BeautifulSoup
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

# ============================================================================
# PINELLAS CITY NAME MAPPING
# ============================================================================

# Map Pinellas County tax district codes to full city names
PINELLAS_CITY_MAP = {
    'SP': 'St. Petersburg',
    'ST PETERSBURG': 'St. Petersburg',
    'ST. PETERSBURG': 'St. Petersburg',
    'CLEARWATER': 'Clearwater',
    'CW': 'Clearwater',
    'LARGO': 'Largo',
    'LA': 'Largo',
    'PINELLAS PARK': 'Pinellas Park',
    'PP': 'Pinellas Park',
    'DUNEDIN': 'Dunedin',
    'TARPON SPRINGS': 'Tarpon Springs',
    'TS': 'Tarpon Springs',
    'SEMINOLE': 'Seminole',
    'KENNETH CITY': 'Kenneth City',
    'GULFPORT': 'Gulfport',
    'MADEIRA BEACH': 'Madeira Beach',
    'REDINGTON BEACH': 'Redington Beach',
    'TREASURE ISLAND': 'Treasure Island',
    'ST PETE BEACH': 'St. Pete Beach',
    'SOUTH PASADENA': 'South Pasadena',
    'BELLEAIR': 'Belleair',
    'BELLEAIR BEACH': 'Belleair Beach',
    'BELLEAIR BLUFFS': 'Belleair Bluffs',
    'INDIAN ROCKS BEACH': 'Indian Rocks Beach',
    'INDIAN SHORES': 'Indian Shores',
    'NORTH REDINGTON BEACH': 'North Redington Beach',
    'OLDSMAR': 'Oldsmar',
    'SAFETY HARBOR': 'Safety Harbor',
    # Unincorporated area codes
    'LFPW': 'Unincorporated Pinellas (Lealman)',
    'LEALMAN': 'Unincorporated Pinellas (Lealman)',
    'UNINCORPORATED': 'Unincorporated Pinellas',
    'COUNTY': 'Unincorporated Pinellas'
}

def expand_city_name(city_abbr):
    """Expand Pinellas city abbreviation to full name."""
    if not city_abbr:
        return 'Unincorporated Pinellas'
    
    city_upper = city_abbr.strip().upper()
    return PINELLAS_CITY_MAP.get(city_upper, city_abbr)

# ============================================================================
# ST. PETERSBURG ZONING/FLU LOOKUP TABLES
# ============================================================================

# St. Petersburg Future Land Use code to description mapping
FLU_DESCRIPTIONS = {
    'CBD': 'Central Business District',
    'CRD': 'Community Redevelopment District',
    'PR-R': 'Planned Redevelopment Residential',
    'PR-MU': 'Planned Redevelopment Mixed-Use',
    'PR-C': 'Planned Redevelopment Commercial',
    'RU': 'Residential Urban',
    'RL': 'Residential Low',
    'RLM': 'Residential Low Medium',
    'RM': 'Residential Medium',
    'RH': 'Residential High',
    'RVH': 'Residential Very High',
    'R/OL': 'Residential/Office Limited',
    'R/OG': 'Residential/Office General',
    'RFH': 'Resort Facilities High',
    'CG': 'Commercial General',
    'IL': 'Industrial Limited',
    'IG': 'Industrial General',
    'P': 'Preservation',
    'R/OS': 'Recreation/Open Space',
    'I': 'Institutional',
    'T/U': 'Transportation/Utility'
}

# St. Petersburg Zoning code to description mapping
ZONING_DESCRIPTIONS = {
    'NT-1': 'NEIGHBORHOOD TRADITIONAL SINGLE-FAMILY-1',
    'NT-2': 'NEIGHBORHOOD TRADITIONAL SINGLE-FAMILY-2',
    'NT-3': 'NEIGHBORHOOD TRADITIONAL SINGLE-FAMILY-3',
    'NT-4': 'NEIGHBORHOOD TRADITIONAL SINGLE-FAMILY-4',
    'NS-1': 'NEIGHBORHOOD SUBURBAN SINGLE-FAMILY-1',
    'NS-2': 'NEIGHBORHOOD SUBURBAN SINGLE-FAMILY-2',
    'NSM-1': 'NEIGHBORHOOD SUBURBAN MULTI-FAMILY-1',
    'NSM-2': 'NEIGHBORHOOD SUBURBAN MULTI-FAMILY-2',
    'NPUD-1': 'NEIGHBORHOOD PLANNED UNIT DEVELOPMENT-1',
    'NPUD-2': 'NEIGHBORHOOD PLANNED UNIT DEVELOPMENT-2',
    'NPUD-3': 'NEIGHBORHOOD PLANNED UNIT DEVELOPMENT-3',
    'NTM-1': 'NEIGHBORHOOD TRADITIONAL MIXED-RESIDENTIAL-1',
    'NMH': 'NEIGHBORHOOD SUBURBAN MOBILE HOME',
    'NSE': 'NEIGHBORHOOD SUBURBAN SINGLE-FAMILY',
    'DC-C': 'DOWNTOWN CENTER-CORE',
    'DC-1': 'DOWNTOWN CENTER-1',
    'DC-2': 'DOWNTOWN CENTER-2',
    'DC-3': 'DOWNTOWN CENTER-3',
    'DC-P': 'DOWNTOWN CENTER-PRESERVATION',
    'CRT-1': 'CORRIDOR RESIDENTIAL TRADITIONAL-1',
    'CRT-2': 'CORRIDOR RESIDENTIAL TRADITIONAL-2',
    'CRS-1': 'CORRIDOR RESIDENTIAL SUBURBAN-1',
    'CRS-2': 'CORRIDOR RESIDENTIAL SUBURBAN-2',
    'CCT-1': 'CORRIDOR COMMERCIAL TRADITIONAL-1',
    'CCT-2': 'CORRIDOR COMMERCIAL TRADITIONAL-2',
    'CCS-1': 'CORRIDOR COMMERCIAL SUBURBAN-1',
    'CCS-2': 'CORRIDOR COMMERCIAL SUBURBAN-2',
    'RC-1': 'RETAIL CENTER-1',
    'RC-2': 'RETAIL CENTER-2',
    'RC-3': 'RETAIL CENTER-3',
    'EC-1': 'EMPLOYMENT CENTERS-1',
    'EC-2': 'EMPLOYMENT CENTERS-2',
    'IS': 'INDUSTRIAL SUBURBAN',
    'IT': 'INDUSTRIAL TRADITIONAL',
    'IC': 'INSTITUTIONAL CENTER',
    'P': 'PRESERVATION',
    'WATER': 'WATER'
}

# NOTE: Unincorporated Pinellas zoning and FLU coded values are now fetched 
# dynamically from layer metadata using fetch_coded_values() function.
# Hardcoded fallbacks kept below in case metadata fetch fails.

# Unincorporated Pinellas County Zoning code to description mapping (FALLBACK)
UNINCORPORATED_ZONING_DESCRIPTIONS = {
    'AL': 'Aquatic Lands',
    'AL-CO': 'Aquatic Lands - Conditional Overlay',
    'AL-W': 'Aquatic Lands - Wellhead Protection Overlay',
    'AL-W-CO': 'Aquatic Lands - Wellhead Protection Overlay - Conditional Overlay',
    'C-1': 'Neighborhood Commercial',
    'C-1-CO': 'Neighborhood Commercial - Conditional Overlay',
    'C-1-H': 'Neighborhood Commerical - Historic District',
    'C-1-W': 'Neighborhood Commercial - Wellhead Protection Overlay',
    'C-1-W-CO': 'Neighborhood Commercial - Wellhead Protection Overlay - Conditional Overlay',
    'CP': 'Commercial Parkway',
    'CP-CO': 'Commercial Parkway - Conditional Overlay',
    'CP-W': 'Commercial Parkway - Wellhead Protection Overlay',
    'CP-W-CO': 'Commercial Parkway - Wellhead Protection Overlay - Conditional Overlay',
    'CR': 'Commercial Recreation',
    'CR-CO': 'Commercial Recreation - Conditional Overlay',
    'CR-W': 'Commercial Recreation - Wellhead Protection Overlay',
    'CR-W-CO': 'Commercial Recreation - Wellhead Protection Overlay - Conditional Overlay',
    'E-1': 'Employment 1',
    'E-1-CO': 'Employment 1 - Conditional Overlay',
    'E-1-W': 'Employment 1 - Wellhead Protection Overlay',
    'E-1-W-CO': 'Employment 1 - Wellhead Protection Overlay - Conditional Overlay',
    'E-2': 'Employment 2',
    'E-2-CO': 'Employment 2 - Conditional Overlay',
    'E-2-W': 'Employment 2 - Wellhead Protection Overlay',
    'E-2-W-CO': 'Employment 2 - Wellhead Protection Overlay - Conditional Overlay',
    'FBC': 'Form Based Code District',
    'FBC-CO': 'Form Based Code District - Conditional Overlay',
    'FBC-W': 'Form Based Code District - Wellhead Protection Overlay',
    'FBC-W-CO': 'Form Based Code District - Wellhead Protection Overlay - Conditional Overlay',
    'FBR': 'Facilities-Based Recreation',
    'FBR-CO': 'Facilities-Based Recreation - Conditional Overlay',
    'FBR-W': 'Facilities-Based Recreation - Wellhead Protection Overlay',
    'FBR-W-CO': 'Facilities-Based Recreation - Wellhead Protection Overlay - Conditional Overlay',
    'GI': 'General Institutional',
    'GI-CO': 'General Institutional - Conditional Overlay',
    'GI-W': 'General Institutional - Wellhead Protection Overlay',
    'GI-W-CO': 'General Institutional - Wellhead Protection Overlay - Conditional Overlay',
    'GO': 'General Office',
    'GO-CO': 'General Office - Conditional Overlay',
    'GO-W': 'General Office - Wellhead Protection Overlay',
    'GO-W-CO': 'General Office - Wellhead Protection Overlay - Conditional Overlay',
    'I': 'Industrial',
    'I-CO': 'Industrial - Conditional Overlay',
    'I-W': 'Industrial - Wellhead Protection Overlay',
    'I-W-CO': 'Industrial - Wellhead Protection Overlay',
    'IPD': 'Industrial Planned Development',
    'IPD-CO': 'Industrial Planned Development - Conditional Overlay',
    'IPD-W': 'Industrial Planned Development - Wellhead Protection Overlay',
    'IPD-W-CO': 'Industrial Planned Development - Wellhead Protection Overlay - Conditional Overlay',
    'LI': 'Limited Institutional',
    'LI-CO': 'Limited Institutional - Conditional Overlay',
    'LI-W': 'Limited Institutional - Wellhead Protection Overlay',
    'LI-W-CO': 'Limited Institutional - Wellhead Protection Overlay - Conditional Overlay',
    'LO': 'Limited Office',
    'LO-CO': 'Limited Office - Conditional Overlay',
    'LO-W': 'Limited Office - Wellhead Protection Overlay',
    'LO-W-CO': 'Limited Office - Wellhead Protection Overlay - Conditional Overlay',
    'MXD': 'Mixed-Use District',
    'MXD-CO': 'Mixed-Use District - Conditional Overlay',
    'MXD-W': 'Mixed-Use District - Wellhead Protection Overlay',
    'MXD-W-CO': 'Mixed-Use District - Wellhead Protection Overlay - Conditional Overlay',
    'OPH-D': 'Old Palm Harbor Downtown',
    'OPH-D-CO': 'Old Palm Harbor Downtown - Conditional Overlay',
    'OPH-D-H': 'Old Palm Harbor Downtown - Historic District',
    'OPH-D-W': 'Old Palm Harbor Downtown - Wellhead Protection Overlay',
    'OPH-D-W-CO': 'Old Palm Harbor Downtown - Wellhead Protection Overlay - Conditional Overlay',
    'P-C': 'Preservation Conservation',
    'P-C-CO': 'Preservation Conservation - Conditional Overlay',
    'P-C-W': 'Preservation Conservation - Wellhead Protection Overlay',
    'P-C-W-CO': 'Preservation Conservation - Wellhead Protection Overlay - Conditional Overlay',
    'P-RM': 'Preservation Resource Management',
    'P-RM-CO': 'Preservation Resource Management - Conditional Overlay',
    'P-RM-W': 'Preservation Resource Management - Wellhead Protection Overlay',
    'P-RM-W-CO': 'Preservation Resource Management - Wellhead Protection Overlay - Conditional Overlay',
    'P.C.AIRPORT': 'PC Airport',
    'P.C.AIRPORT-CO': 'PC Airport - Conditional Overlay',
    'P.C.AIRPORT-W': 'PC Airport - Wellhead Protection Overlay',
    'P.C.AIRPORT-W-CO': 'PC Airport- Wellhead Protection Overlay - Conditional Overlay',
    'R-1': 'Single Family Residential (9,500 SF Min)',
    'R-1-CO': 'Single Family Residential (9,500 SF Min) - Conditional Overlay',
    'R-1-W': 'Single Family Residential (9,500 SF Min) - Wellhead Protection Overlay',
    'R-1-W-CO': 'Single Family Residential (9,500 SF Min) - Wellhead Protection Overlay - Conditional Overlay',
    'R-2': 'Single Family Residential (7,500 SF Min)',
    'R-2-CO': 'Single Family Residential (7,500 SF Min) - Conditional Overlay',
    'R-2-W': 'Single Family Residential (7,500 SF Min) - Wellhead Protection Overlay',
    'R-2-W-CO': 'Single Family Residential (7,500 SF Min) - Wellhead Protection Overlay - Conditional Overlay',
    'R-3': 'Single Family Residential (6,000 SF Min)',
    'R-3-CO': 'Single Family Residential (6,000 SF Min) - Conditional Overlay',
    'R-3-H': 'Single Family Residential (6,000 SF Min) - Historic District',
    'R-3-W': 'Single Family Residential (6,000 SF Min) - Wellhead Protection Overlay',
    'R-3-W-CO': 'Single Family Residential (6,000 SF Min) - Wellhead Protection Overlay - Conditional Overlay',
    'R-4': 'One, Two and Three Family Residential',
    'R-4-CO': 'One, Two and Three Family Residential - Conditional Overlay',
    'R-4-W': 'One, Two and Three Family Residential - Wellhead Protection Overlay',
    'R-4-W-CO': 'One, Two and Three Family Residential - Wellhead Protection Overlay - Conditional Overlay',
    'R-5': 'Urban Residential District',
    'R-5-CO': 'Urban Residential District - Conditional Overlay',
    'R-5-W': 'Urban Residential District - Wellhead Protection Overlay',
    'R-5-W-CO': 'Urban Residential District - Wellhead Protection Overlay - Conditional Overlay',
    'R-A': 'Residential Agriculture',
    'R-A-CO': 'Residential Agriculture - Conditional Overlay',
    'R-A-W': 'Residential Agriculture - Wellhead Protection Overlay',
    'R-A-W-CO': 'Residential Agriculture - Wellhead Protection Overlay - Conditional Overlay',
    'R-E': 'Residential Estate',
    'R-E-C-T': 'Residential Estate - Transient Accommodation Overlay',
    'R-E-CO': 'Residential Estate - Conditional Overlay',
    'R-E-W': 'Residential Estate - Wellhead Protection Overlay',
    'R-E-W-CO': 'Residential Estate - Wellhead Protection Overlay - Conditional Overlay',
    'R-R': 'Rural Residential',
    'R-R-CO': 'Rural Residential - Conditional Overlay',
    'R-R-H': 'Rural Residential - Historic District',
    'R-R-W': 'Rural Residnetial - Wellhead Protection Overlay',
    'R-R-W-CO': 'Rural Residential - Wellhead Protection Overlay - Conditional Overlay',
    'RBR': 'Resource-Based Recreation',
    'RBR-CO': 'Resource-Based Recreation - Conditional Overlay',
    'RBR-W': 'Resource-Based Recreation - Wellhead Protection Overlay',
    'RBR-W-CO': 'Resource-Based Recreation - Wellhead Protection Overlay - Conditional Overlay',
    'RM': 'Multi-Family Residential (see FLUM for density)',
    'RM-CO': 'Multi-Family Residential (see FLUM for density) - Conditional Overlay',
    'RM-W': 'Multi-Family Residential (see FLUM for density) - Wellhead Protection Overlay',
    'RM-W-CO': 'Multi-Family Residential (see FLUM for density) - Wellhead Protection Overlay - Conditional Overlay',
    'RMH': 'Residential Mobile/Manufactured Home',
    'RMH-CO': 'Residential Mobile/Manufactured Home - Conditional Overlay',
    'RMH-W': 'Residential Mobile/Manufactured Home - Wellhead Protection Overlay',
    'RMH-W-CO': 'Residential Mobile/Manufactured Home - Wellhead Protection Overlay - Conditional Overlay',
    'RPD': 'Residential Planned Development (see FLUM for density)',
    'RPD-CO': 'Residential Planned Development (see FLUM for density) - Conditional Overlay',
    'RPD-W': 'Residential Planned Developlment (see FLUM for density) - Wellhead Protection Overlay',
    'RPD-W-CO': 'Residential Planned Development (see FLUM for density) - Wellhead Protection Overlay - Conditional Overlay',
    'UZ': 'Unknown Zoning',
    'UZ-CO': 'Unknown Zoning - Conditional Overlay',
    'UZ-W': 'Unknown Zoning - Wellhead Protection Overlay',
    'UZ-W-CO': 'Unknown Zoning - Wellhead Protection Overlay - Conditional Overlay',
    'OPH-D-W': 'Old Palm Harbor Downtown - Wellhead Protection Overlay',
    'C-T': 'Transient Accommodation Overlay',
    'HPO': 'Historic Preservation Overlay',
    'E-1-C-T': 'Employment 1 - Transient Accommodation Overlay',
    'C-2': 'General Commercial and Services',
    'DPH-FBC': 'Downtown Palm Harbor Form Based Code',
    'C-2-C-T': 'General Commercial and Services Transient Accommodations Overlay',
    'L-FBC': 'Lealman - Form Based Code',
    'C-2-CO': 'General Commercial and Services - Conditional Overlay',
    'C-2-H': 'General Commercial and Services - Historic District',
    'C-2-W': 'General Commercial and Services - Wellhead Protection Overlay',
    'C-2-W-CO': 'General Commercial and Services - Wellhead Proteciton Overlay - Conditional Overlay',
}

# Unincorporated Pinellas County Future Land Use code to description mapping (FALLBACK)
UNINCORPORATED_FLU_DESCRIPTIONS = {
    'RR': 'Residential Rural',
    'RE': 'Residential Estate',
    'RS': 'Residential Suburban',
    'RL': 'Residential Low',
    'RU': 'Residential Urban',
    'RLM': 'Residential Low Medium',
    'RM': 'Residential Medium',
    'RH': 'Residential High',
    'PR-I': 'Planned Redevelopment - Industrial',
    'RFO': 'Resort Facilities',
    'PR-C': 'Planned Redevelopment - Commercial',
    'NO-DES': 'No Designation',
    'CN': 'Commercial Neighborhood',
    'CG': 'Commercial General',
    'CR': 'Commercial Recreation',
    'IL': 'Industrial Limited',
    'IG': 'Industrial General',
    'P': 'Preservation',
    'PR-MU': 'Planned Redevelopment - Mixed Use',
    'I': 'Institutional',
    'PR-R': 'Planned Redevelopment - Residential',
    'TU': 'Transportation/Utilities',
    'ROR': 'Residential/Office/Retail',
    'ROL': 'Residential/Office/Limited',
    'ROG': 'Residential/Office/General',
    'ROS': 'Recreation/Open Space',
    'PSP': 'Public/Semi-Public',
    'RVH': 'Residential Very High',
    'RFM': 'Resort Facilities Medium',
    'RFH': 'Resort Facilities High',
    'CRD': 'Community Redevelopment Dist',
    'CBD': 'Central Business District',
    'CL': 'Commercial Limited',
    'RFO-P': 'Resort Facilities Overlay/Perm',
    'RFO-T': 'Resort Facilities Overlay/Temp',
    'WDF': 'Water Drainage Feature',
    'WF': 'Water Feature',
    'WATER': 'WATER',
    'ROAD': 'ROAD',
    'P-RM': 'Preservation - Resource Management',
    'MUNI': 'MUNICIPALITY',
    'NO-D-W': 'No Designation Uninc Water',
    'MUNI-W': 'Municipal Open Water',
    'CRD-AC': 'Community Redevelop-Activity Ctr',
    'AC': 'AC',
    'AC-P': 'AC-P',
    'RM-12.5': 'RM-12.5',
    'TU-O': 'Transportation/Utility Overlay',
    'E': 'Employment',
    'AC-N': 'Activity Center - Neighborhood',
    'AC-C': 'Activity Center - Community',
    'AC-M': 'Activity Center - Major',
    'MUC-P': 'Mixed Use Corridor - Primary',
    'MUC-S': 'Mixed Use Corridor - Secondary',
    'MUC-P-C': 'Mixed Use Corridor - Primary - Commerce',
    'MUC-SU-NP': 'Mixed Use Corridor - Supporting - Neighborhood Park',
    'MUC-SU-LT': 'Mixed Use Corridor - Supporting - Local Trade',
}

# ============================================================================
# AUTOMATED CODED VALUE EXTRACTION FROM ARCGIS LAYERS
# ============================================================================

@st.cache_data(ttl=86400, show_spinner=False)
def fetch_coded_values(layer_url, field_name):
    """
    Automatically fetch coded value domain from an ArcGIS layer.
    
    Args:
        layer_url: Base layer URL (e.g., "https://...MapServer/1")
        field_name: Field with coded values (e.g., "ZONEDESC")
    
    Returns:
        dict: {code: description} mapping, or empty dict if not found
    """
    session = get_resilient_session()
    
    try:
        # Fetch layer metadata
        metadata_url = f"{layer_url}?f=json"
        response = session.get(metadata_url, timeout=15)
        response.raise_for_status()
        metadata = response.json()
        
        # Find the field with coded values
        for field in metadata.get('fields', []):
            if field.get('name') == field_name:
                domain = field.get('domain')
                if domain and domain.get('type') == 'codedValue':
                    # Extract code -> name mappings
                    coded_values = domain.get('codedValues', [])
                    return {item['code']: item['name'] for item in coded_values}
        
        return {}
    
    except Exception as e:
        # Return empty dict on error - calling code should handle
        return {}

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def validate_parcel_id(parcel_id: str):
    """Validate parcel/folio ID input."""
    if not parcel_id:
        return False, "Parcel ID cannot be empty"
    
    if len(parcel_id) > 30:
        return False, "Parcel ID must be 30 characters or less"
    
    # Allowlist: only alphanumeric, dashes, spaces, periods
    if not re.match(r'^[A-Za-z0-9\-\s\.]+$', parcel_id):
        return False, "Invalid characters in parcel ID"
    
    return True, ""

def sanitize_for_sql(value: str) -> str:
    """Sanitize string for use in SQL WHERE clause."""
    return value.strip().replace("'", "''")

def strip_dor_code(land_use_text):
    """Remove Florida DOR code prefix from land use descriptions."""
    if not land_use_text:
        return ''
    
    text = land_use_text.strip()
    
    if text and text[0].isdigit():
        parts = text.split(' ', 1)
        if len(parts) > 1:
            return parts[1].strip()
    
    return text

# ============================================================================
# HTTP SESSION WITH RETRY LOGIC
# ============================================================================

@st.cache_resource
def get_resilient_session():
    """Create HTTP session with automatic retry logic."""
    session = requests.Session()
    
    retry_strategy = Retry(
        total=3,
        backoff_factor=1.0,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["GET", "POST"]
    )
    
    adapter = HTTPAdapter(max_retries=retry_strategy)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    
    return session

# ============================================================================
# PCPAO API LOOKUP (THE KEY FUNCTION!)
# ============================================================================

def scrape_pinellas_property(parcel_id):
    """
    Query Pinellas County Property Appraiser searchProperty API.
    This is the backend API that the PCPAO website uses.
    """
    session = get_resilient_session()
    
    url = "https://www.pcpao.gov/dal/quicksearch/searchProperty"
    
    # Normalize Pinellas parcel ID format
    normalized_parcel = parcel_id.strip()
    
    # If no dashes and 18 digits, add dashes in Pinellas format
    if '-' not in normalized_parcel and len(normalized_parcel) == 18:
        # Format: XX-XX-XX-XXXXX-XXX-XXXX
        normalized_parcel = f"{normalized_parcel[0:2]}-{normalized_parcel[2:4]}-{normalized_parcel[4:6]}-{normalized_parcel[6:11]}-{normalized_parcel[11:14]}-{normalized_parcel[14:18]}"
    
    # Build the POST data - mimics the DataTables request format
    payload = {
        'draw': '1',
        'start': '0',
        'length': '10',
        'search[value]': '',
        'search[regex]': 'false',
        'input': normalized_parcel,
        'searchsort': 'parcel_number',
        'url': 'https://www.pcpao.gov'
    }
    
    # Add column definitions (required by DataTables API)
    for i in range(11):
        payload[f'columns[{i}][data]'] = str(i)
        payload[f'columns[{i}][name]'] = ''
        payload[f'columns[{i}][searchable]'] = 'true'
        payload[f'columns[{i}][orderable]'] = 'true' if i >= 2 else 'false'
        payload[f'columns[{i}][search][value]'] = ''
        payload[f'columns[{i}][search][regex]'] = 'false'
    
    try:
        response = session.post(url, data=payload, timeout=15)
        response.raise_for_status()
        
        data = response.json()
        
        # Check if we got results
        if data.get('recordsTotal', 0) == 0:
            return {'success': False, 'error': 'Parcel not found in PCPAO database'}
        
        # Parse the HTML response within the JSON
        if not data.get('data') or len(data['data']) == 0:
            return {'success': False, 'error': 'No property data returned'}
        
        # Get first result
        result_row = data['data'][0]
        
        # Extract data from HTML snippets
        # Column 2: Owner name
        owner_html = result_row[2] if len(result_row) > 2 else ''
        owner_soup = BeautifulSoup(owner_html, 'lxml')
        owner = owner_soup.get_text(strip=True)
        
        # Column 5: Address
        address_html = result_row[5] if len(result_row) > 5 else ''
        address_soup = BeautifulSoup(address_html, 'lxml')
        address = address_soup.get_text(strip=True)
        
        # Column 6: Current Tax District (this IS the city, but may be abbreviated)
        tax_dist_html = result_row[6] if len(result_row) > 6 else ''
        tax_dist_soup = BeautifulSoup(tax_dist_html, 'lxml')
        tax_district = tax_dist_soup.get_text(strip=True)
        
        # Expand abbreviated city name (e.g., "SP" -> "St. Petersburg")
        city = expand_city_name(tax_district)
        
        # Column 7: Property Use / DOR Code
        use_html = result_row[7] if len(result_row) > 7 else ''
        use_soup = BeautifulSoup(use_html, 'lxml')
        property_use = use_soup.get_text(strip=True)
        
        # Column 8: Legal Description
        legal_html = result_row[8] if len(result_row) > 8 else ''
        legal_soup = BeautifulSoup(legal_html, 'lxml')
        legal_desc = legal_soup.get_text(strip=True)
        
        # Get acreage from detail page
        sqft = None
        acres = None
        zip_code = None
        
        try:
            # Strap transformation: swap first and third segments
            parts = normalized_parcel.split('-')
            if len(parts) == 6:
                parts[0], parts[2] = parts[2], parts[0]
                strap = ''.join(parts)
            else:
                strap = normalized_parcel.replace('-', '')
            
            # Build detail URL
            detail_url = (
                f"https://www.pcpao.gov/property-details?"
                f"s={strap}&"
                f"input={normalized_parcel}&"
                f"search_option=parcel_number"
            )
            
            # Fetch and parse
            html = session.get(detail_url, timeout=30).text
            soup = BeautifulSoup(html, "html.parser")
            text = soup.get_text(" ", strip=True)
            
            # Match pattern: "Land Area: ≅ 59,560 sf | ≅ 1.36 acres"
            m = re.search(r"Land Area:\s*≅\s*([\d,]+)\s*sf\s*\|\s*≅\s*([\d.]+)\s*acres", text)
            if m:
                sqft = int(m.group(1).replace(",", ""))
                acres = float(m.group(2))
            
            # Extract ZIP code from detail page (format: "FL 33703" or "FL33703")
            zip_match = re.search(r'FL\s*(\d{5})', text)
            if zip_match:
                zip_code = zip_match.group(1)
        except Exception:
            pass  # If detail page fails, sqft, acres, and zip_code remain None
        
        return {
            'success': True,
            'address': address,
            'city': city,
            'zip': zip_code or '',
            'owner': owner,
            'land_use': strip_dor_code(property_use),
            'zoning': 'Contact City/County for zoning info',
            'site_area_sqft': f"{sqft:,}" if sqft else None,
            'site_area_acres': f"{acres:.2f}" if acres else None,
            'legal_description': legal_desc,
            'error': None
        }
    
    except Exception as e:
        return {'success': False, 'error': f'Error querying PCPAO API: {str(e)}'}

# ============================================================================
# ST. PETERSBURG ZONING LAYER LOOKUP
# ============================================================================

def is_unincorporated(city_name):
    """Check if city is unincorporated Pinellas."""
    if not city_name:
        return True
    
    unincorporated_indicators = [
        'UNINCORPORATED',
        'LFPW',  # Lealman area
        'LEALMAN',
        'COUNTY',
        'PINELLAS COUNTY'
    ]
    
    city_upper = city_name.upper()
    return any(indicator in city_upper for indicator in unincorporated_indicators)

def lookup_unincorporated_zoning(address):
    """
    Lookup zoning and FLU for unincorporated Pinellas County areas.
    Uses PublicWebGIS/Landuse_Zoning/MapServer with automated coded value extraction.
    
    Returns: dict with zoning_code, zoning_description, future_land_use, future_land_use_description
    """
    if not address:
        return {'success': False, 'error': 'Address required for zoning lookup'}
    
    session = get_resilient_session()
    
    try:
        # Fetch coded value domains dynamically (cached for 24 hours)
        zoning_lookup = fetch_coded_values(
            "https://egis.pinellas.gov/gis/rest/services/PublicWebGIS/Landuse_Zoning/MapServer/1",
            "ZONEDESC"
        )
        flu_lookup = fetch_coded_values(
            "https://egis.pinellas.gov/gis/rest/services/PublicWebGIS/Landuse_Zoning/MapServer/0",
            "LANDUSEDESC"
        )
        
        # Fallback to hardcoded if fetch failed
        if not zoning_lookup:
            zoning_lookup = UNINCORPORATED_ZONING_DESCRIPTIONS
        if not flu_lookup:
            flu_lookup = UNINCORPORATED_FLU_DESCRIPTIONS
        
        # Step 1: Geocode the address
        search_url = "https://geocode.arcgis.com/arcgis/rest/services/World/GeocodeServer/findAddressCandidates"
        geocode_params = {
            'SingleLine': f"{address}, Pinellas County, FL",
            'f': 'json',
            'outFields': '*'
        }
        
        geocode_response = session.get(search_url, params=geocode_params, timeout=15)
        geocode_data = geocode_response.json()
        
        if not geocode_data.get('candidates'):
            return {'success': False, 'error': 'Could not geocode address'}
        
        # Get coordinates
        location = geocode_data['candidates'][0]['location']
        x, y = location['x'], location['y']
        
        # Step 2: Query Zoning layer (Layer 1 - Zoning - Unincorporated)
        zoning_url = "https://egis.pinellas.gov/gis/rest/services/PublicWebGIS/Landuse_Zoning/MapServer/1/query"
        zoning_params = {
            'geometry': f"{x},{y}",
            'geometryType': 'esriGeometryPoint',
            'inSR': '4326',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'ZONEDESC',
            'returnGeometry': 'false',
            'f': 'json'
        }
        
        zoning_response = session.get(zoning_url, params=zoning_params, timeout=15)
        zoning_data = zoning_response.json()
        
        zoning_code = ''
        zoning_desc = ''
        if zoning_data.get('features'):
            zoning_attrs = zoning_data['features'][0]['attributes']
            zoning_code = zoning_attrs.get('ZONEDESC', '')  # This returns the CODE
            # Look up the description from fetched or fallback dictionary
            zoning_desc = zoning_lookup.get(zoning_code, '')
        
        # Step 3: Query Future Land Use layer (Layer 0)
        flu_url = "https://egis.pinellas.gov/gis/rest/services/PublicWebGIS/Landuse_Zoning/MapServer/0/query"
        flu_params = {
            'geometry': f"{x},{y}",
            'geometryType': 'esriGeometryPoint',
            'inSR': '4326',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'LANDUSECODE,LANDUSEDESC',
            'returnGeometry': 'false',
            'f': 'json'
        }
        
        flu_response = session.get(flu_url, params=flu_params, timeout=15)
        flu_data = flu_response.json()
        
        flu_code = ''
        flu_desc = ''
        if flu_data.get('features'):
            flu_attrs = flu_data['features'][0]['attributes']
            # Get the code from either field (they both return the code)
            flu_code = flu_attrs.get('LANDUSECODE') or flu_attrs.get('LANDUSEDESC', '')
            # Look up the description from fetched or fallback dictionary
            flu_desc = flu_lookup.get(flu_code, '')
        
        return {
            'success': True,
            'zoning_code': zoning_code,
            'zoning_description': zoning_desc,
            'future_land_use': flu_code,
            'future_land_use_description': flu_desc
        }
        
    except Exception as e:
        return {'success': False, 'error': f'Unincorporated zoning lookup error: {str(e)}'}

def lookup_clearwater_zoning(address):
    """
    Lookup zoning and FLU for Clearwater properties.
    Uses Clearwater's own GIS services.
    
    Returns: dict with zoning_code, zoning_description, future_land_use, future_land_use_description
    """
    if not address:
        return {'success': False, 'error': 'Address required for zoning lookup'}
    
    session = get_resilient_session()
    
    try:
        # Step 1: Geocode the address
        search_url = "https://geocode.arcgis.com/arcgis/rest/services/World/GeocodeServer/findAddressCandidates"
        geocode_params = {
            'SingleLine': f"{address}, Clearwater, FL",
            'f': 'json',
            'outFields': '*'
        }
        
        geocode_response = session.get(search_url, params=geocode_params, timeout=15)
        geocode_data = geocode_response.json()
        
        if not geocode_data.get('candidates'):
            return {'success': False, 'error': 'Could not geocode address'}
        
        # Get coordinates
        location = geocode_data['candidates'][0]['location']
        x, y = location['x'], location['y']
        
        # Step 2: Query Zoning layer
        zoning_url = "https://gis.myclearwater.com/arcgis/rest/services/ArcGISMapServices/Zoning_WGS84/MapServer/1/query"
        zoning_params = {
            'geometry': f"{x},{y}",
            'geometryType': 'esriGeometryPoint',
            'inSR': '4326',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'ZONING,ZONING_DESC',
            'returnGeometry': 'false',
            'f': 'json'
        }
        
        zoning_response = session.get(zoning_url, params=zoning_params, timeout=15)
        zoning_data = zoning_response.json()
        
        zoning_code = ''
        zoning_desc = ''
        if zoning_data.get('features'):
            zoning_attrs = zoning_data['features'][0]['attributes']
            zoning_code = zoning_attrs.get('ZONING', '')
            zoning_desc = zoning_attrs.get('ZONING_DESC', '')
        
        # Step 3: Query Future Land Use layer (Layer 0 confirmed)
        # Fetch coded values for FLU
        flu_lookup = fetch_coded_values(
            "https://gis.myclearwater.com/arcgis/rest/services/ArcGISMapServices/FLU_w_PPC_Colors_WGS84/MapServer/0",
            "LU"
        )
        
        flu_url = "https://gis.myclearwater.com/arcgis/rest/services/ArcGISMapServices/FLU_w_PPC_Colors_WGS84/MapServer/0/query"
        flu_params = {
            'geometry': f"{x},{y}",
            'geometryType': 'esriGeometryPoint',
            'inSR': '4326',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'LU',  # The actual field name!
            'returnGeometry': 'false',
            'f': 'json'
        }
        
        flu_response = session.get(flu_url, params=flu_params, timeout=15)
        flu_data = flu_response.json()
        
        flu_code = ''
        flu_desc = ''
        if flu_data.get('features'):
            flu_attrs = flu_data['features'][0]['attributes']
            flu_code = flu_attrs.get('LU', '')
            # Look up description from coded values
            flu_desc = flu_lookup.get(flu_code, '')
        
        return {
            'success': True,
            'zoning_code': zoning_code,
            'zoning_description': zoning_desc,
            'future_land_use': flu_code,
            'future_land_use_description': flu_desc
        }
        
    except Exception as e:
        return {'success': False, 'error': f'Clearwater zoning lookup error: {str(e)}'}

def lookup_largo_zoning(address, parcel_data=None):
    """
    Lookup zoning and FLU for Largo properties.
    Largo uses Future Land Use classification instead of traditional zoning.
    Queries Largo's parcel layer for Countywide_Plan_Map_Category_1.
    
    Args:
        address: Property address
        parcel_data: Not used (kept for compatibility)
    
    Returns: dict with zoning_code, future_land_use, descriptions
    """
    if not address:
        return {'success': False, 'error': 'Address required'}
    
    session = get_resilient_session()
    
    try:
        # Step 1: Geocode the address
        search_url = "https://geocode.arcgis.com/arcgis/rest/services/World/GeocodeServer/findAddressCandidates"
        geocode_params = {
            'SingleLine': f"{address}, Largo, FL",
            'f': 'json',
            'outFields': '*'
        }
        
        geocode_response = session.get(search_url, params=geocode_params, timeout=15)
        geocode_data = geocode_response.json()
        
        if not geocode_data.get('candidates'):
            return {'success': False, 'error': 'Could not geocode address'}
        
        # Get coordinates
        location = geocode_data['candidates'][0]['location']
        x, y = location['x'], location['y']
        
        # Step 2: Query Largo parcel layer for Countywide Plan category
        parcel_url = "https://maps.largo.com/arcgis/rest/services/Largo_GIS_Viewer_Map/MapServer/247/query"
        parcel_params = {
            'geometry': f"{x},{y}",
            'geometryType': 'esriGeometryPoint',
            'inSR': '4326',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'Countywide_Plan_Map_Category_1',
            'returnGeometry': 'false',
            'f': 'json'
        }
        
        parcel_response = session.get(parcel_url, params=parcel_params, timeout=15)
        parcel_data_result = parcel_response.json()
        
        flu_value = ''
        if parcel_data_result.get('features'):
            attrs = parcel_data_result['features'][0]['attributes']
            flu_value = attrs.get('Countywide_Plan_Map_Category_1', '')
        
        if flu_value:
            return {
                'success': True,
                'zoning_code': flu_value,  # Largo uses FLU as zoning
                'zoning_description': None,  # Already combined in flu_value
                'future_land_use': flu_value,  # Same as zoning for Largo
                'future_land_use_description': None  # Already combined
            }
        else:
            return {
                'success': False,
                'error': 'Could not retrieve Countywide Plan category from Largo parcel layer'
            }
        
    except Exception as e:
        return {'success': False, 'error': f'Largo lookup error: {str(e)}'}

def lookup_pinellas_zoning(city_name, address, parcel_data=None):
    """
    Router function: Lookup zoning for Pinellas County based on city.
    Routes to appropriate lookup function based on jurisdiction.
    
    Args:
        city_name: City name (e.g., "St. Petersburg", "Lealman", "Clearwater", "Largo")
        address: Property address (e.g., "200 CENTRAL AVE")
        parcel_data: Optional parcel data from PCPAO lookup (used for Largo)
        
    Returns:
        dict with zoning_code, zoning_description, future_land_use, future_land_use_description
    """
    if not address:
        return {'success': False, 'error': 'Address required for zoning lookup'}
    
    # Route based on jurisdiction
    if 'St. Petersburg' in city_name or 'St Petersburg' in city_name:
        # St. Petersburg has its own GIS layers
        return lookup_stpete_zoning(address)
    elif 'Clearwater' in city_name:
        # Clearwater has its own GIS services
        return lookup_clearwater_zoning(address)
    elif 'Largo' in city_name:
        # Largo uses FLU instead of traditional zoning
        return lookup_largo_zoning(address, parcel_data)
    elif is_unincorporated(city_name):
        # Unincorporated areas use Pinellas County layers
        return lookup_unincorporated_zoning(address)
    else:
        # Other municipalities - not yet implemented
        return {
            'success': True,
            'zoning_code': 'Contact City for zoning',
            'zoning_description': None,
            'future_land_use': None,
            'future_land_use_description': None,
            'note': f'City-specific zoning data for {city_name} not yet implemented'
        }

# ============================================================================
# FLORIDA DOR LAND USE CODES (for Hillsborough)
# ============================================================================

DOR_LAND_USE_CODES = {
    '0000': 'Vacant Residential', '0001': 'Single Family', '0002': 'Mobile Homes',
    '0003': 'Multi-Family (10+ units)', '0004': 'Condominiums', '0005': 'Cooperatives',
    '0006': 'Retirement Homes', '0007': 'Miscellaneous Residential',
    '0008': 'Multi-Family (Less than 10 units)', '0009': 'Residential Common Elements/Areas',
    '0010': 'Vacant Commercial', '0011': 'Stores, One Story', '0012': 'Mixed Use Store/Office',
    '0013': 'Department Stores', '0014': 'Supermarkets', '0015': 'Regional Shopping Centers',
    '0016': 'Community Shopping Centers', '0017': 'Office Buildings, One Story',
    '0018': 'Office Buildings, Multi-Story', '0019': 'Professional Service Buildings',
    '0020': 'Airports, Terminals, Marinas', '0021': 'Restaurants, Cafeterias',
    '0022': 'Drive-In Restaurants', '0023': 'Financial Institutions',
    '0024': 'Insurance Company Offices', '0025': 'Repair Service Shops',
    '0026': 'Service Stations', '0027': 'Auto Sales, Auto Repair',
    '0028': 'Parking Lots, Mobile Home Parks', '0029': 'Wholesale Outlets',
    '0030': 'Florists, Greenhouses', '0031': 'Drive-In Theaters',
    '0032': 'Enclosed Theaters', '0033': 'Nightclubs, Bars',
    '0034': 'Bowling Alleys, Skating Rinks', '0035': 'Tourist Attractions',
    '0036': 'Camps', '0037': 'Race Tracks', '0038': 'Golf Courses', '0039': 'Hotels, Motels',
    '0040': 'Vacant Industrial', '0041': 'Light Manufacturing', '0042': 'Heavy Industrial',
    '0043': 'Lumber Yards, Sawmills', '0044': 'Packing Plants', '0045': 'Canneries, Bottlers',
    '0046': 'Other Food Processing', '0047': 'Mineral Processing', '0048': 'Warehousing',
    '0049': 'Open Storage', '0050': 'Improved Agricultural', '0051': 'Cropland Class I',
    '0052': 'Cropland Class II', '0053': 'Cropland Class III', '0054': 'Timberland',
    '0055': 'Timberland', '0056': 'Timberland', '0057': 'Timberland', '0058': 'Timberland',
    '0059': 'Woods, Native Pasture', '0060': 'Grazing Land Class I',
    '0061': 'Grazing Land Class II', '0062': 'Grazing Land Class III',
    '0063': 'Grazing Land Class IV', '0064': 'Grazing Land Class V',
    '0065': 'Grazing Land Class VI', '0066': 'Orchard Groves, Citrus',
    '0067': 'Poultry, Bees, Fish', '0068': 'Dairies, Feed Lots',
    '0069': 'Ornamentals, Misc Agricultural', '0070': 'Vacant Institutional',
    '0071': 'Churches', '0072': 'Private Schools/Colleges', '0073': 'Private Hospitals',
    '0074': 'Homes for the Aged', '0075': 'Non-Profit/Charitable',
    '0076': 'Mortuaries, Cemeteries', '0077': 'Clubs, Lodges',
    '0078': 'Sanitariums, Convalescent Homes', '0079': 'Cultural Organizations',
    '0080': 'Vacant Governmental', '0081': 'Military', '0082': 'Forests, Parks, Recreational',
    '0083': 'Public Schools', '0084': 'Colleges (Government)', '0085': 'Hospitals (Government)',
    '0086': 'County Buildings', '0087': 'State Buildings', '0088': 'Federal Buildings',
    '0089': 'Municipal Buildings', '0090': 'Leasehold Interests', '0091': 'Utility Buildings',
    '0092': 'Mining Lands', '0093': 'Petroleum and Gas', '0094': 'Telephone Exchange Buildings',
    '0095': 'Convenience Stores', '0096': 'Sewage Disposal, Waste Land',
    '0097': 'Outdoor Recreational', '0098': 'Centrally Assessed',
    '0099': 'Acreage Not Zoned Agricultural',
    # 4-digit codes
    '8400': 'SCHOOLS/COLLEGE', '8300': 'PUBLIC SCHOOLS', '8500': 'HOSPITALS (GOVERNMENT)',
    # 2-digit versions for codes that come without leading zeros
    '86': 'County Buildings', '87': 'State Buildings', '88': 'Federal Buildings',
    '89': 'Municipal Buildings', '83': 'Public Schools', '84': 'Colleges, Public', 
    '85': 'Hospitals, Public', '91': 'Utility Buildings', '93': 'Petroleum and Gas',
    '94': 'Telephone Exchange Buildings', '95': 'Convenience Stores', '98': 'Centrally Assessed',
}

# Pasco County MAF Use Codes (from search.pascopa.com/codes.aspx?type=016)
PASCO_USE_CODES = {
    '00': 'Unimproved', '01': 'Single Family Residential', '02': 'Mobile Home',
    '03': 'Multi Family (5 or more Units per Building)', '04': 'Condominium',
    '05': 'Multi Story Retirement Apartments', '07': 'Single Family Villas',
    '08': 'Multi Family (4 or less Units per Building)', '11': 'Retail Stores (One Story)',
    '12': 'Stores / Office SFR', '13': 'Department Stores, Discount Department Stores',
    '14': 'Grocery Store, Market', '16': 'Shopping Center Local',
    '17': 'Offices (One Story)', '18': 'Offices (Multi Story)',
    '19': 'Offices Professional or Medical', '20': 'Airports',
    '21': 'Restaurants (Full Service)', '22': 'Restaurants (Fast Food)',
    '23': 'Financial Institutions', '25': 'Service Shops',
    '26': 'Service Stations, Old Style', '27': 'Vehicle Sales / Dealer / Repair',
    '29': 'Wholesale Manufacturing', '30': 'Florist, Greenhouses',
    '31': 'Theaters, Drive In', '32': 'Theaters, Enclosed',
    '33': 'Bars, Lounges, Night Clubs', '34': 'Bowling Alleys, Arenas',
    '35': 'Tourist Attraction', '36': 'Camps', '37': 'Race Tracks',
    '38': 'Golf Courses', '39': 'Hotels Motels', '41': 'Light Manufacturing',
    '42': 'Heavy Industrial', '43': 'Lumber Yards', '44': 'Packing Plants',
    '45': 'Breweries, Wineries, Etc.', '46': 'Metal Bldg Misc',
    '47': 'Mineral Processing', '48': 'Warehouses,Block', '49': 'Open Space',
    '50': 'Improved Rural Homesite', '51': 'Cropland Class I',
    '52': 'Cropland Class II', '53': 'Cropland Class II',
    '54': 'Timber, Site Index I', '55': 'Timber, Site Index II',
    '56': 'Timber, Site Index III', '57': 'Timber, Site Index IV',
    '58': 'Timber, Site Index V', '59': 'Unclassified Timberland',
    '60': 'Grazing Land Class I', '61': 'Grazing Land Class II',
    '62': 'Grazing Land Class III', '63': 'Grazing Land Class IV',
    '64': 'Grazing Land Class V', '65': 'Grazing Land Class VI',
    '66': 'Orchards, Groves', '67': 'Poultry, Bees, etc.',
    '68': 'Daries, Feed Lots', '69': 'Ornamentals',
    '70': 'Vacant Institutional', '71': 'Churches',
    '72': 'Schools, Colleges, Private', '73': 'Hospitals, Private',
    '74': 'Nursing Homes (Typical Skilled Nursing Facility)',
    '75': 'Orphanage', '76': 'Mortuary, Funeral Home',
    '77': 'Clubs, Lodges, Halls', '78': 'Inpatient/Outpatient Clinics',
    '79': 'Warehouses, Mini Storage', '80': 'Metal Buildings',
    '81': 'Military', '82': 'Forest, Park, Etc.', '83': 'Schools, Public',
    '84': 'Colleges, Public', '85': 'Hospitals, Public',
    '86': 'County Buildings', '87': 'State Buildings',
    '88': 'Federal Buildings', '89': 'Municipal Buildings',
    '91': 'Utility Buildings', '93': 'Petroleum and Gas',
    '94': 'Telephone Exchange Buildings', '95': 'Convenience Stores',
    '98': 'Centrally Assessed',
}

def get_pasco_use_description(code):
    """Convert Pasco County MAF use code to description."""
    if not code:
        return ''
    code_str = str(code).strip()
    # If already text, return as-is
    if any(c.isalpha() for c in code_str):
        return code_str
    # Strip leading zeros for lookup (086 -> 86)
    code_str = code_str.lstrip('0') or '00'  # Keep '00' as '00', not empty
    # Look up in Pasco codes
    return PASCO_USE_CODES.get(code_str, code_str)

def get_land_use_description(code_or_desc):
    """Convert DOR land use code to description. If already text, return as-is."""
    if not code_or_desc:
        return ''
    code_str = str(code_or_desc).strip()
    if any(c.isalpha() for c in code_str):
        return code_str
    
    # Remove decimal point (84.0 -> 84)
    if '.' in code_str:
        code_str = code_str.split('.')[0]
    
    if code_str in DOR_LAND_USE_CODES:
        return DOR_LAND_USE_CODES[code_str]
    code_padded = code_str.zfill(4)
    if code_padded in DOR_LAND_USE_CODES:
        return DOR_LAND_USE_CODES[code_padded]
    return code_str

HILLSBOROUGH_FLU_CODES = {
    'A/M': 'Agricultural/Mining - 1/20', 'A-1/10': 'Agricultural - 1/10',
    'AR-1/5': 'Agricultural/Rural - 1/5', 'RES-1': 'Residential - 1',
    'RES-2': 'Residential - 2', 'RES-4': 'Residential - 4',
    'RES-6': 'Residential - 6', 'RES-9': 'Residential - 9',
    'RES-12': 'Residential - 12', 'RES-16': 'Residential - 16',
    'RES-20': 'Residential - 20', 'RES-35': 'Residential - 35',
    'LI': 'Light Industrial', 'HI': 'Heavy Industrial',
    'P/Q-P': 'Public/Quasi-Public', 'P/QP': 'Public/Quasi-Public',
    'E': 'Environmentally Sensitive Area',
}

# Pasco County Zoning Descriptions
PASCO_ZONING_DESCRIPTIONS = {
    'MPUD': 'Master Planned Unit Development',
    'MPUD-100': 'Master Planned Unit Development',
    'PD': 'Planned Development',
    'AG': 'Agricultural',
    'AG-1': 'Agricultural-1',
    'RR': 'Rural Residential',
    'RE': 'Residential Estate',
    'RS': 'Residential Suburban',
    'R4': 'Residential-4',
    'RM': 'Residential Medium',
    'RH': 'Residential High',
    'MH': 'Mobile Home',
    'CN': 'Commercial Neighborhood',
    'CG': 'Commercial General',
    'IL': 'Industrial Limited',
    'IG': 'Industrial General',
    'P/SP': 'Public/Semi-Public',
}

def get_hillsborough_flu_description(code):
    """Convert Hillsborough FLU code to description."""
    if not code:
        return ''
    return HILLSBOROUGH_FLU_CODES.get(str(code).strip(), code)

# ============================================================================
# HILLSBOROUGH COUNTY FUNCTIONS  
# ============================================================================

def lookup_hillsborough_parcel(folio):
    """Lookup Hillsborough parcel by folio from SWFWMD service."""
    try:
        parcel_url = "https://www25.swfwmd.state.fl.us/arcgis12/rest/services/BaseVector/parcel_search/MapServer/7/query"
        
        folio_formats = [folio, folio.replace('.', ''), folio.split('.')[0] if '.' in folio else folio]
        
        parcel_data = None
        for folio_format in folio_formats:
            parcel_params = {
                'where': f"FOLIONUM = '{folio_format}'",
                'outFields': '*',
                'returnGeometry': 'true',
                'outSR': '4326',
                'f': 'json'
            }
            parcel_resp = requests.get(parcel_url, params=parcel_params, timeout=15)
            data = parcel_resp.json()
            if data.get('features'):
                parcel_data = data
                break
        
        if not parcel_data or not parcel_data.get('features'):
            return {'success': False, 'error': f'Folio not found'}
        
        attrs = parcel_data['features'][0]['attributes']
        geom = parcel_data['features'][0].get('geometry')
        
        acres = attrs.get('ACRES') or attrs.get('AREANO')
        acres_str = f"{float(acres):.2f}" if acres and acres not in [None, 'None', ''] else ''
        
        dor_code = attrs.get('DORUSECODE') or attrs.get('DOR4CODE')
        land_use_desc = get_land_use_description(dor_code) if dor_code else attrs.get('PARUSEDESC', '')
        
        return {
            'success': True,
            'address': attrs.get('SITEADD', attrs.get('SITUSADD1', '')),
            'city': attrs.get('SCITY', 'Tampa'),
            'zip': attrs.get('SZIP', ''),
            'owner': attrs.get('OWNNAME', attrs.get('OWNERNAME', '')),
            'land_use': land_use_desc,
            'site_area_acres': acres_str,
            'site_area_sqft': '',
            'zoning': attrs.get('ZONING', ''),
            'geometry': geom
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}

def lookup_hillsborough_zoning_flu(address, geometry=None):
    """Lookup Hillsborough zoning/FLU from official county GIS layers."""
    try:
        # Use provided geometry if available, otherwise geocode
        if geometry and geometry.get('rings'):
            # Calculate centroid from polygon
            ring = geometry['rings'][0]
            x_coords = [pt[0] for pt in ring]
            y_coords = [pt[1] for pt in ring]
            x = sum(x_coords) / len(x_coords)
            y = sum(y_coords) / len(y_coords)
        else:
            # Geocode the address
            geocode_url = "https://geocode.arcgis.com/arcgis/rest/services/World/GeocodeServer/findAddressCandidates"
            geocode_params = {'SingleLine': address + ", Hillsborough County, FL", 'f': 'json', 'outSR': '4326', 'maxLocations': 1}
            
            geocode_resp = requests.get(geocode_url, params=geocode_params, timeout=10)
            geocode_data = geocode_resp.json()
            
            if not geocode_data.get('candidates'):
                return {'success': False, 'error': 'Could not geocode address'}
            
            location = geocode_data['candidates'][0]['location']
            x, y = location['x'], location['y']
        
        result = {'success': True}
        
        # Create envelope (bounding box) around the point
        buffer = 0.0001
        xmin, ymin = x - buffer, y - buffer
        xmax, ymax = x + buffer, y + buffer
        
        # Query Official Hillsborough County Zoning Atlas
        zoning_url = "https://maps.hillsboroughcounty.org/arcgis/rest/services/DSD_Viewer_Services/DSD_Viewer_Zoning_Regulatory/MapServer/0/query"
        zoning_params = {
            'where': '1=1',
            'geometry': f'{xmin},{ymin},{xmax},{ymax}',
            'geometryType': 'esriGeometryEnvelope',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'NZONE,NZONE_DESC,CATEGORY',
            'returnGeometry': 'false',
            'f': 'json',
            'inSR': '4326'
        }
        
        zoning_resp = requests.get(zoning_url, params=zoning_params, timeout=10)
        zoning_data = zoning_resp.json()
        
        if zoning_data.get('features'):
            attrs = zoning_data['features'][0]['attributes']
            result['zoning_code'] = attrs.get('NZONE', '')
            result['zoning_description'] = attrs.get('NZONE_DESC', '')
        
        # Query Future Land Use layer
        flu_url = "https://maps.hillsboroughcounty.org/arcgis/rest/services/DSD_Viewer_Services/DSD_Viewer_Planning/MapServer/1/query"
        flu_params = {
            'where': '1=1',
            'geometry': f'{xmin},{ymin},{xmax},{ymax}',
            'geometryType': 'esriGeometryEnvelope',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'FLUE',
            'returnGeometry': 'false',
            'f': 'json',
            'inSR': '4326'
        }
        
        flu_resp = requests.get(flu_url, params=flu_params, timeout=10)
        flu_data = flu_resp.json()
        
        if flu_data.get('features'):
            flu_code = flu_data['features'][0]['attributes'].get('FLUE', '')
            result['future_land_use'] = get_hillsborough_flu_description(flu_code)
        
        return result
    except Exception as e:
        return {'success': False, 'error': str(e)}

# ============================================================================
# PASCO COUNTY FUNCTIONS  
# ============================================================================

def lookup_pasco_parcel(parcel_id):
    """Lookup Pasco parcel by parcel ID from SWFWMD service."""
    try:
        parcel_url = "https://www25.swfwmd.state.fl.us/arcgis12/rest/services/BaseVector/parcel_search/MapServer/12/query"
        
        # Try with and without dashes
        parcel_formats = [parcel_id, parcel_id.replace('-', '')]
        
        parcel_data = None
        for parcel_format in parcel_formats:
            parcel_params = {
                'where': f"PARCELID = '{parcel_format}'",
                'outFields': '*',
                'returnGeometry': 'true',
                'outSR': '4326',
                'f': 'json'
            }
            parcel_resp = requests.get(parcel_url, params=parcel_params, timeout=15)
            data = parcel_resp.json()
            if data.get('features'):
                parcel_data = data
                break
        
        if not parcel_data or not parcel_data.get('features'):
            return {'success': False, 'error': f'Parcel not found'}
        
        attrs = parcel_data['features'][0]['attributes']
        geom = parcel_data['features'][0].get('geometry')
        
        # Try site address first, fallback to mailing address
        address = attrs.get('SITEADD') or attrs.get('SITUSADD1') or attrs.get('MAILADD') or ''
        city = attrs.get('SCITY') or attrs.get('MCITY') or ''
        zip_code = attrs.get('SZIP') or attrs.get('MZIP') or ''
        
        acres = attrs.get('ACRES') or attrs.get('AREANO')
        acres_str = f"{float(acres):.2f}" if acres and acres not in [None, 'None', ''] else ''
        
        # Pasco: PARUSEDESC usually has text already (like "Vacant Commercial")
        # Only convert if it's a numeric code (like "86")
        parusedesc = attrs.get('PARUSEDESC', '')
        if parusedesc and parusedesc.strip().replace('.', '').isdigit():
            # It's a numeric code - convert using MAF codes
            land_use_desc = get_pasco_use_description(parusedesc.strip())
        else:
            # It's already text - use it directly
            land_use_desc = parusedesc
        
        return {
            'success': True,
            'address': address,
            'city': city,
            'zip': zip_code,
            'owner': attrs.get('OWNNAME', attrs.get('OWNERNAME', '')),
            'land_use': land_use_desc,
            'site_area_acres': acres_str,
            'site_area_sqft': '',
            'zoning': '',
            'flu': '',
            'geometry': geom
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}

def lookup_pasco_zoning_flu(address, geometry=None):
    """Lookup Pasco zoning/FLU from county GIS layers."""
    try:
        # Use provided geometry if available, otherwise geocode
        if geometry and geometry.get('rings'):
            ring = geometry['rings'][0]
            x_coords = [pt[0] for pt in ring]
            y_coords = [pt[1] for pt in ring]
            x = sum(x_coords) / len(x_coords)
            y = sum(y_coords) / len(y_coords)
        else:
            geocode_url = "https://geocode.arcgis.com/arcgis/rest/services/World/GeocodeServer/findAddressCandidates"
            geocode_params = {'SingleLine': address + ", Pasco County, FL", 'f': 'json', 'outSR': '4326', 'maxLocations': 1}
            
            geocode_resp = requests.get(geocode_url, params=geocode_params, timeout=10)
            geocode_data = geocode_resp.json()
            
            if not geocode_data.get('candidates'):
                return {'success': False, 'error': 'Could not geocode address'}
            
            location = geocode_data['candidates'][0]['location']
            x, y = location['x'], location['y']
        
        result = {'success': True}
        
        # Create envelope around the point
        buffer = 0.0001
        xmin, ymin = x - buffer, y - buffer
        xmax, ymax = x + buffer, y + buffer
        
        # Query Pasco County Zoning (Layer 1)
        zoning_url = "https://mapping.pascopa.com/arcgis/rest/services/Land_Use/MapServer/1/query"
        zoning_params = {
            'where': '1=1',
            'geometry': f'{xmin},{ymin},{xmax},{ymax}',
            'geometryType': 'esriGeometryEnvelope',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': '*',
            'returnGeometry': 'false',
            'f': 'json',
            'inSR': '4326'
        }
        
        zoning_resp = requests.get(zoning_url, params=zoning_params, timeout=10)
        zoning_data = zoning_resp.json()
        
        
        if zoning_data.get('features'):
            attrs = zoning_data['features'][0]['attributes']
            
            # Use ZN_TYPE for zoning code (R4, MPUD, etc.)
            zoning_code = attrs.get('ZN_TYPE', '')
            zoning_desc = PASCO_ZONING_DESCRIPTIONS.get(zoning_code, '')
            
            result['zoning_code'] = zoning_code
            result['zoning_description'] = zoning_desc
        
        # ALWAYS query Future Land Use (Layer 0)
        flu_url = "https://mapping.pascopa.com/arcgis/rest/services/Land_Use/MapServer/0/query"
        flu_params = {
            'where': '1=1',
            'geometry': f'{xmin},{ymin},{xmax},{ymax}',
            'geometryType': 'esriGeometryEnvelope',
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': '*',
            'returnGeometry': 'false',
            'f': 'json',
            'inSR': '4326'
        }
        
        flu_resp = requests.get(flu_url, params=flu_params, timeout=10)
        flu_data = flu_resp.json()
        
        
        if flu_data.get('features'):
            attrs = flu_data['features'][0]['attributes']
            # Correct field names: FLU_CODE and DESCRIPTION
            flu_code = attrs.get('FLU_CODE', '')
            flu_desc = attrs.get('DESCRIPTION', '')
            # Combine code and description
            if flu_code and flu_desc:
                result['future_land_use'] = f"{flu_code} - {flu_desc}"
            elif flu_desc:
                result['future_land_use'] = flu_desc
            elif flu_code:
                result['future_land_use'] = flu_code
        
        return result
    except Exception as e:
        return {'success': False, 'error': str(e)}

def lookup_stpete_zoning(address):
    """
    Lookup zoning for St. Petersburg properties using St. Pete GIS layers.
    (Renamed from original lookup_pinellas_zoning St. Pete section)
    
    Args:
        address: Property address (e.g., "200 CENTRAL AVE")
        
    Returns:
        dict with zoning_code, zoning_description, future_land_use, future_land_use_description
    """
    session = get_resilient_session()
    
    try:
        # Step 1: Geocode the address to get coordinates
        search_url = "https://geocode.arcgis.com/arcgis/rest/services/World/GeocodeServer/findAddressCandidates"
        geocode_params = {
            'SingleLine': f"{address}, St. Petersburg, FL",
            'f': 'json',
            'outFields': '*'
        }
        
        geocode_response = session.get(search_url, params=geocode_params, timeout=15)
        geocode_data = geocode_response.json()
        
        if not geocode_data.get('candidates'):
            return {'success': False, 'error': 'Could not geocode address'}
        
        # Get coordinates from first candidate
        location = geocode_data['candidates'][0]['location']
        x, y = location['x'], location['y']
        
        # Step 2: Query zoning layer with coordinates (spatial query)
        zoning_url = "https://egis.stpete.org/arcgis/rest/services/ServicesDSD/Zoning/MapServer/2/query"
        zoning_params = {
            'geometry': f"{x},{y}",
            'geometryType': 'esriGeometryPoint',
            'inSR': '4326',  # WGS84 from geocoder
            'spatialRel': 'esriSpatialRelIntersects',
            'outFields': 'ZONECLASS,ZONEDESC',
            'returnGeometry': 'false',
            'f': 'json'
        }
        
        zoning_response = session.get(zoning_url, params=zoning_params, timeout=15)
        zoning_data = zoning_response.json()
        
        if zoning_data.get('features'):
            attrs = zoning_data['features'][0]['attributes']
            zoning_code = attrs.get('ZONECLASS', '')
            zoning_desc = ZONING_DESCRIPTIONS.get(zoning_code, attrs.get('ZONEDESC', ''))
            
            # Step 3: Query Future Land Use layer with same coordinates
            flu_url = "https://egis.stpete.org/arcgis/rest/services/ServicesDSD/Zoning/MapServer/4/query"
            flu_params = {
                'geometry': f"{x},{y}",
                'geometryType': 'esriGeometryPoint',
                'inSR': '4326',
                'spatialRel': 'esriSpatialRelIntersects',
                'outFields': '*',
                'returnGeometry': 'false',
                'f': 'json'
            }
            
            flu_response = session.get(flu_url, params=flu_params, timeout=15)
            flu_data = flu_response.json()
            flu_code = ''
            flu_desc = ''
            if flu_data.get('features'):
                flu_attrs = flu_data['features'][0].get('attributes', {})
                flu_code = flu_attrs.get('LANDUSECODE', '')
                flu_desc = FLU_DESCRIPTIONS.get(flu_code, '')
            
            return {
                'success': True,
                'zoning_code': zoning_code,
                'zoning_description': zoning_desc,
                'future_land_use': flu_code,
                'future_land_use_description': flu_desc
            }
        else:
            return {'success': False, 'error': 'No zoning found at address location'}
            
    except Exception as e:
        return {'success': False, 'error': f'St. Pete zoning lookup error: {str(e)}'}
    
    # Other cities (not St. Petersburg)
    return {
        'success': True,
        'zoning_code': 'Contact City/County for zoning',
        'zoning_description': None,
        'future_land_use': None,
        'future_land_use_description': None,
        'note': 'City-specific zoning data not available via API'
    }

# ============================================================================
# STREAMLIT UI
# ============================================================================


# ============================================================================
# DOCUMENT GENERATION FUNCTIONS
# ============================================================================

# DOCUMENT GENERATION FUNCTIONS
# ============================================================================

def set_cell_background(cell, color_hex):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing_shd = tcPr.find(qn('w:shd'))
    if existing_shd is not None:
        tcPr.remove(existing_shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def set_cell_margins(cell, top=20, bottom=20, start=40, end=40):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = OxmlElement('w:tcMar')
    for margin_name, value in [('top', top), ('bottom', bottom), ('start', start), ('end', end)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(value))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)


def remove_table_borders(table):
    """Remove all borders from table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def create_header(section):
    """Create Kimley-Horn header."""
    header = section.header
    header.is_linked_to_previous = False
    
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    header_table.columns[0].width = Inches(5.0)
    header_table.columns[1].width = Inches(1.5)
    
    tbl = header_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    
    logo_cell = header_table.cell(0, 0)
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_para.clear()
    
    run1 = logo_para.add_run("Kimley")
    run1.font.size = Pt(28)
    run1.font.bold = False
    run1.font.color.rgb = RGBColor(88, 89, 91)
    run1.font.name = 'Arial Narrow'
    
    run2 = logo_para.add_run("»")
    run2.font.size = Pt(28)
    run2.font.bold = False
    run2.font.color.rgb = RGBColor(88, 89, 91)
    run2.font.name = 'Arial Narrow'
    
    run3 = logo_para.add_run("Horn")
    run3.font.size = Pt(28)
    run3.font.bold = False
    run3.font.color.rgb = RGBColor(166, 25, 46)
    run3.font.name = 'Arial Narrow'
    
    page_cell = header_table.cell(0, 1)
    page_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    page_para = page_cell.paragraphs[0]
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = page_para.add_run('Page ')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def create_footer(section):
    """Create Kimley-Horn footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    
    widths = [Inches(1.1), Inches(0.01), Inches(4.23), Inches(0.01), Inches(0.96)]
    colors = ['5F5F5F', None, 'A20C33', None, 'A20C33']
    texts = ['kimley-horn.com', '', '200 Central Avenue Suite 600 St. Petersburg, FL 33701', '', '(727) 822-5150']
    
    table = footer.add_table(rows=1, cols=5, width=sum(widths))
    table.allow_autofit = False
    remove_table_borders(table)
    
    row = table.rows[0]
    row.height = Inches(0.22)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    
    for idx, cell in enumerate(row.cells):
        table.columns[idx].width = widths[idx]
        cell.width = widths[idx]
        
        if colors[idx]:
            set_cell_background(cell, colors[idx])
        
        # Only set margins for colored cells, not gap cells
        if colors[idx]:
            set_cell_margins(cell, top=20, bottom=20, start=40, end=40)
        else:
            # Gap cells get zero margins for precise 0.01" spacing
            set_cell_margins(cell, top=0, bottom=0, start=0, end=0)
        
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
        if texts[idx]:
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.clear()
            
            run = para.add_run(texts[idx])
            run.font.name = 'Arial'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(255, 255, 255)


def add_opening_section(doc, client_info, project_info):
    """Add opening section with proper letterhead format matching DS template."""
    
    # Date at top
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = date_para.add_run(project_info['date'])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    date_para.paragraph_format.space_after = Pt(0)
    date_para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Client Name (short name)
    para = doc.add_paragraph()
    run = para.add_run(client_info['name'])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Legal Entity Name (exact per SunBiz)
    para = doc.add_paragraph()
    run = para.add_run(client_info.get('legal_entity', client_info['name']))
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Entity Address Line 1
    para = doc.add_paragraph()
    run = para.add_run(client_info['address1'])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Entity Address Line 2
    para = doc.add_paragraph()
    run = para.add_run(client_info['address2'])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # "Re:" line with project info
    para = doc.add_paragraph()
    run = para.add_run('Re:\t' + project_info["name"])
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Project address (if available)
    if project_info.get('address'):
        para = doc.add_paragraph()
        run = para.add_run('\t' + project_info['address'])
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
    
    # City, State
    para = doc.add_paragraph()
    run = para.add_run('\t' + project_info.get('city_state_zip', project_info['city'] + ', Florida'))
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # "Dear Client:" salutation
    para = doc.add_paragraph()
    run = para.add_run(f'Dear {client_info["contact"]}:')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Opening paragraph
    para = doc.add_paragraph()
    opening_text = f'Kimley-Horn and Associates, Inc. ("Kimley-Horn") is pleased to submit this letter agreement (the "Agreement") to {client_info.get("legal_entity", client_info["name"])} ("the Client") for professional consulting services for the above referenced project.'
    run = para.add_run(opening_text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()


def add_project_understanding(doc, project_info, assumptions):
    """Add Project Understanding section with user's description and assumptions."""
    
    # Section heading - BOLD + CENTERED (no underline per template)
    para = doc.add_paragraph()
    run = para.add_run('PROJECT UNDERSTANDING')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # User's project description/understanding - JUSTIFIED
    if project_info.get('description'):
        para = doc.add_paragraph()
        run = para.add_run(project_info['description'])
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Assumptions intro
    para = doc.add_paragraph()
    run = para.add_run('Kimley-Horn understands the following in preparing this proposal:')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Assumptions as bullet points
    for assumption in assumptions:
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(assumption)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Closing statement
    para = doc.add_paragraph()
    run = para.add_run('If any of these assumptions are not correct, then the scope and fee will change. Based on the above understanding, Kimley-Horn proposes the following scope of services:')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()


def add_scope_of_services(doc, selected_tasks, permits, included_additional_services=None):
    """Add Scope of Services section."""
    
    if included_additional_services is None:
        included_additional_services = []
    
    # Section heading - BOLD + CENTERED (no underline per template)
    para = doc.add_paragraph()
    run = para.add_run('Scope of Services')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    run = para.add_run('Kimley-Horn will provide the services specifically set forth below.')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    sub_section_keywords = ['cover sheet', 'general notes', 'utility plan', 'site layout', 'site plan',
                           'grading plan', 'drainage plan', 'paving', 'erosion control',
                           'detail', 'existing conditions', 'demolition', 'stormwater pollution',
                           'civil details', 'construction specifications']
    
    for task_num in sorted(selected_tasks.keys()):
        task = selected_tasks[task_num]
        
        # Special handling for Task 310 - Construction Phase Services
        if task_num == '310' and 'hours' in task:
            hours = task['hours']
            descriptions = []
            for desc in TASK_DESCRIPTIONS[task_num]:
                # Replace hour placeholders
                desc = desc.replace('{shop_drawing_hours}', str(hours['shop_drawing']))
                desc = desc.replace('{rfi_hours}', str(hours['rfi']))
                desc = desc.replace('{oac_meetings}', str(hours['oac_meetings']))
                desc = desc.replace('{site_visits}', str(hours['site_visits']))
                desc = desc.replace('{record_drawing_hours}', str(hours['record_drawing']))
                desc = desc.replace('{total_hours}', str(hours['total']))
                descriptions.append(desc)
        # Special handling for Task 150 - Civil Permitting
        elif task_num == '150' and permits:
            descriptions = [
                "Kimley-Horn will prepare and submit on the Client's behalf the civil construction documents to the following agencies:"
            ]
            # Add permits as separate items for bullet point formatting
            descriptions.extend(permits)
            descriptions.extend([
                "Includes development of a Stormwater Ownership and Maintenance (O&M) Plan",
                "Kimley-Horn will respond to a maximum of three (3) requests for information during the agency review process for obtaining the above permits.",
                "**BOLD:**Permit fees and impact fees are not included. Kimley-Horn does not guarantee the issuance of permits or approvals."
            ])
        else:
            descriptions = TASK_DESCRIPTIONS[task_num]
        
        # Add blank line before task heading
        doc.add_paragraph()
        
        # Task heading - BOLD + UNDERLINED
        para = doc.add_paragraph()
        run = para.add_run(f'Task {task_num} – {task["name"].replace("Civil ", "")}')
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.underline = True
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        
        # Removed: doc.add_paragraph() here - was creating double spacing
        
        permit_list_started = False
        
        for desc in descriptions:
            # Check if this is a permit item (short, from permits list)
            is_permit_bullet = (task_num == '150' and desc in permits)
            
            # Start bullet list after intro
            if task_num == '150' and 'following agencies:' in desc:
                para = doc.add_paragraph()
                run = para.add_run(desc)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0
                doc.add_paragraph()
                permit_list_started = True
                continue
            
            if is_permit_bullet:
                # Permit as bullet point - INDENTED
                para = doc.add_paragraph(style='List Bullet')
                para.paragraph_format.left_indent = Inches(0.25)  # Indent for clarity
                run = para.add_run(desc)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0
                continue
            
            # End permit bullets, back to regular paragraphs
            if permit_list_started and not is_permit_bullet:
                permit_list_started = False
                doc.add_paragraph()
            
            # Check if bold text
            is_bold_para = desc.startswith('**BOLD:**')
            
            # Check if sub-section heading
            is_subsection = (len(desc) < 100 and 
                           any(kw in desc.lower() for kw in sub_section_keywords) and
                           not desc.endswith('.') and
                           not desc.startswith('**NOTE:**') and
                           not desc.startswith('**BOLD:**'))
            
            # Add one blank line before subsection headings OR notes
            if is_subsection or desc.startswith('**NOTE:**'):
                doc.add_paragraph()
            
            para = doc.add_paragraph()
            
            # Check if it's a note
            is_note = desc.startswith('**NOTE:**')
            
            if is_note:
                run = para.add_run(desc.replace('**NOTE:**', 'Note:'))
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.italic = True
            elif is_bold_para:
                run = para.add_run(desc.replace('**BOLD:**', ''))
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
            else:
                run = para.add_run(desc)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                
                if is_subsection:
                    run.font.underline = True
            
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            
            # Add blank line after Note paragraphs
            if is_note:
                doc.add_paragraph()
    
    # Add included additional services (if any) as additional tasks
    if included_additional_services:
        doc.add_paragraph()
        
        # Add a note about additional services being included
        para = doc.add_paragraph()
        run = para.add_run('Additional Services Included in This Proposal:')
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.0
        
        # List the included additional services as bullets
        for service in included_additional_services:
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(service)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
        
        doc.add_paragraph()


def add_scope_table(doc, selected_tasks, included_additional_services_with_fees=None):
    """Add Scope of Work table with optional additional services."""
    
    if included_additional_services_with_fees is None:
        included_additional_services_with_fees = {}
    
    # Calculate total including additional services
    total_fee = sum(task['fee'] for task in selected_tasks.values())
    total_fee += sum(included_additional_services_with_fees.values())
    
    # Calculate number of rows: header + regular tasks + additional services + total
    num_additional = len(included_additional_services_with_fees)
    num_rows = len(selected_tasks) + num_additional + 2
    
    table = doc.add_table(rows=num_rows, cols=4)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Task Number & Name'
    header_cells[1].text = 'Task Number & Name'
    header_cells[2].text = 'Fee'
    header_cells[3].text = 'Type'
    
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = 'Arial'
    
    # Add regular tasks
    row_idx = 1
    for task_num, task in sorted(selected_tasks.items()):
        row = table.rows[row_idx]
        row.cells[0].text = task_num
        row.cells[1].text = task['name']
        row.cells[2].text = f'$ {task["fee"]:,}'
        row.cells[3].text = task['type']
        
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.name = 'Arial'
        
        row_idx += 1
    
    # Add included additional services as AS-1, AS-2, etc.
    as_num = 1
    for service_name, service_fee in included_additional_services_with_fees.items():
        row = table.rows[row_idx]
        row.cells[0].text = f'AS-{as_num}'
        row.cells[1].text = service_name
        row.cells[2].text = f'$ {service_fee:,}'
        row.cells[3].text = 'Hourly, Not-to-Exceed'
        
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.name = 'Arial'
        
        row_idx += 1
        as_num += 1
    
    # Total row
    total_row = table.rows[-1]
    total_row.cells[0].text = 'Total'
    total_row.cells[1].text = 'Total'
    total_row.cells[2].text = f'$ {total_fee:,}'
    total_row.cells[3].text = f'$ {total_fee:,}'
    
    for cell in total_row.cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = 'Arial'
    
    doc.add_paragraph()
    
    task_list = ', '.join(sorted(selected_tasks.keys()))
    para = doc.add_paragraph()
    run = para.add_run(f'Kimley-Horn will perform the services in Tasks {task_list} on a labor fee plus expense basis with the maximum labor fee shown above.')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0


def add_additional_services_section(doc, excluded_additional_services):
    """Add ADDITIONAL SERVICES section per template - lists what's NOT included."""
    
    # Only add this section if there are excluded services to list
    if not excluded_additional_services:
        return
    
    # Section heading - BOLD + CENTERED + DARK RED (no underline per template)
    para = doc.add_paragraph()
    run = para.add_run('ADDITIONAL SERVICES')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(162, 12, 51)  # A20C33 - matches footer center column
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Opening paragraph
    para = doc.add_paragraph()
    text = 'Based on the information of which we are aware, we have prepared a proposal that we believe to be comprehensive. In the event that an unforeseen issue(s) should arise, we remain available to provide additional services, as requested by you, on the basis of our hourly rates or an agreed upon lump sum amount. Potential services not addressed in this proposal are:'
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    # List of excluded services (what's NOT included)
    for service in excluded_additional_services:
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(service)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()


def add_information_provided_section(doc):
    """Add INFORMATION PROVIDED BY THE CLIENT section per template."""
    
    # Section heading - BOLD + CENTERED + DARK RED (no underline per template)
    para = doc.add_paragraph()
    run = para.add_run('INFORMATION PROVIDED BY THE CLIENT')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(162, 12, 51)  # A20C33 - matches footer center column
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Opening paragraph
    para = doc.add_paragraph()
    text = 'The following information, upon which the consultant may rely, will be provided to Kimley-Horn by the Client or its representative:'
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    # List of items client must provide
    client_provided_items = [
        "Building footprints in AutoCAD format",
        "Boundary, Topographic, and Tree survey in AutoCAD format",
        "All Application/Permit Fees",
        "Geotechnical survey",
        "Building elevations"
    ]
    
    for item in client_provided_items:
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(item)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()


def add_client_responsibilities_section(doc):
    """Add CLIENT RESPONSIBILITIES section per template."""
    
    # Section heading - BOLD + CENTERED + DARK RED (no underline per template)
    para = doc.add_paragraph()
    run = para.add_run('CLIENT RESPONSIBILITIES')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(162, 12, 51)  # A20C33 - matches footer center column
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Opening paragraph
    para = doc.add_paragraph()
    text = 'In addition to other responsibilities set out in this Agreement, the Client shall:'
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    # List of client responsibilities
    client_responsibilities = [
        "Provide access to the project site(s) or other land which Kimley-Horn is conduct any field work in a timely manner.",
        "Provide prompt notice whenever it observes or otherwise becomes aware of any development that affects the scope or timing of Kimley-Horn's performance.",
        "Examine and provide comments and/or decisions with respect to any Kimley-Horn interim or final deliverables within a period mutually agreed upon."
    ]
    
    for responsibility in client_responsibilities:
        para = doc.add_paragraph(style='List Bullet')
        run = para.add_run(responsibility)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()


def add_schedule_section(doc):
    """Add SCHEDULE section per template."""
    
    # Section heading - BOLD + CENTERED + DARK RED (no underline per template)
    para = doc.add_paragraph()
    run = para.add_run('SCHEDULE')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(162, 12, 51)  # A20C33 - matches footer center column
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Schedule paragraph
    para = doc.add_paragraph()
    text = 'Consultant shall provide the services described in the above scope as expeditiously as practical to meet a mutually agreed upon schedule. This Agreement is made in anticipation of continuous permitting conditions and orderly progress through completion of the services. The mutually agreed upon schedule shall be extended as necessary for delays or suspensions due to circumstances that the Consultant does not control.'
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()


def add_fee_and_billing_section(doc, selected_tasks):
    """Add FEE AND BILLING section per template - shows all applicable fee types."""
    
    # Section heading - BOLD + CENTERED + DARK RED (no underline per template)
    para = doc.add_paragraph()
    run = para.add_run('FEE AND BILLING')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(162, 12, 51)  # A20C33 - matches footer center column
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Group tasks by fee type
    lump_sum_tasks = []
    hourly_tasks = []
    hourly_nte_tasks = []
    
    for task_num, task in selected_tasks.items():
        fee_type = task.get('type', 'Hourly, Not-to-Exceed')
        if fee_type == 'Lump Sum':
            lump_sum_tasks.append(task_num)
        elif fee_type == 'Hourly':
            hourly_tasks.append(task_num)
        else:  # Hourly, Not-to-Exceed
            hourly_nte_tasks.append(task_num)
    
    # Add paragraph for each fee type that's used
    
    # Lump Sum paragraph
    if lump_sum_tasks:
        para = doc.add_paragraph()
        task_list = ', '.join(sorted(lump_sum_tasks))
        text = f'Kimley-Horn will perform the Services in Tasks {task_list} for the total lump sum labor fee shown in the table above.'
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.0
    
    # Hourly (Cost Plus) paragraph
    if hourly_tasks:
        para = doc.add_paragraph()
        task_list = ', '.join(sorted(hourly_tasks))
        text = f'Kimley-Horn will perform the Services in Tasks {task_list} on a labor fee plus expense basis. Labor fee will be billed on an hourly basis according to Kimley-Horn\'s then-current rates.'
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.0
    
    # Hourly Not-to-Exceed (Cost Plus Max) paragraph
    if hourly_nte_tasks:
        para = doc.add_paragraph()
        task_list = ', '.join(sorted(hourly_nte_tasks))
        text = f'Kimley-Horn will perform the services in Tasks {task_list} on a labor fee plus expense basis with the maximum labor fee shown in the table above. Labor fee will be billed on an hourly basis according to Kimley-Horn\'s then-current rates.'
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.0
        
        # Add Not-to-exceed clause
        para = doc.add_paragraph()
        text = 'Kimley-Horn will not exceed the total maximum labor fee shown without authorization from the Client. However, Kimley-Horn reserves the right to reallocate amounts among tasks as necessary.'
        run = para.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.0
    
    # Direct reimbursable expenses paragraph (always included)
    para = doc.add_paragraph()
    text = 'Direct reimbursable expenses such as express delivery services, air travel, and other direct expenses are not included in the above fee table and will be billed at 1.15 times cost. A percentage of labor fee will be added to each invoice to cover certain other expenses such as telecommunications, in-house reproduction, postage, supplies, project related computer time, and local mileage. Administrative time related to the project will be billed hourly. All permitting, application, and similar project fees will be paid directly by the Client. Should the Client request Kimley-Horn to advance any such project fees on the Client\'s behalf, an invoice for such fees, with a fifteen percent (15%) markup, will be immediately issued to and paid by the Client.'
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    # Invoicing and payment terms
    para = doc.add_paragraph()
    text = 'Fees and expenses will be invoiced monthly based, as applicable, upon the percentage of services completed or actual services performed, plus expenses incurred as of the invoice date. Payment will be due within 25 days of the date of the invoice and should include the invoice number and Kimley-Horn project number.'
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()


def add_closure_section(doc, client_info, invoice_info):
    """Add CLOSURE section with signature blocks per template."""
    
    # Section heading - BOLD + CENTERED + DARK RED (no underline per template)
    para = doc.add_paragraph()
    run = para.add_run('CLOSURE')
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(162, 12, 51)  # A20C33 - matches footer center column
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Standard Provisions reference
    para = doc.add_paragraph()
    text = f'In addition to the matters set forth herein, our Agreement shall include and be subject to, and only to, the terms and conditions in the attached Standard Provisions, which are incorporated by reference. As used in the Standard Provisions, the term "the Consultant" shall refer to Kimley-Horn and Associates, Inc., and the term "the Client" shall refer to {client_info.get("legal_entity", client_info.get("name", ""))}.'.strip()
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Electronic invoicing paragraph
    para = doc.add_paragraph()
    text = 'Kimley-Horn, in an effort to expedite invoices and reduce paper waste, offers its clients the option to receive electronic invoices. These invoices come via email in an Adobe PDF format. Please select a billing method from the choices below:'
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    # Invoice email line 1
    para = doc.add_paragraph()
    invoice_email = invoice_info.get('email', '_' * 60)
    run = para.add_run(f'____ Please email all invoices to {invoice_email}')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.line_spacing = 1.0
    
    # Invoice CC email line 2
    para = doc.add_paragraph()
    invoice_cc_email = invoice_info.get('cc_email', '_' * 60)
    if invoice_cc_email:
        run = para.add_run(f'____ Please copy {invoice_cc_email}')
    else:
        run = para.add_run(f'____ Please copy {"_" * 60}')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Closing paragraph with conditional retainer language
    para = doc.add_paragraph()
    
    # Build text based on whether retainer is required
    if invoice_info.get('use_retainer') and invoice_info.get('retainer_amount', 0) > 0:
        retainer_amt = invoice_info['retainer_amount']
        text = f'To proceed with the services, please have an authorized person sign this Agreement below and return to us with a retainer of ${retainer_amt:,}. We will commence services only after we have received a fully-executed agreement and a retainer in the amount of ${retainer_amt:,}. Fees and times stated in this Agreement are valid for sixty (60) days after the date of this letter.'
    else:
        text = 'To proceed with the services, please have an authorized person sign this Agreement below and return to us. We will commence services only after we have received a fully-executed agreement. Fees and times stated in this Agreement are valid for sixty (60) days after the date of this letter.'
    
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # "We appreciate..." paragraph
    para = doc.add_paragraph()
    text = 'We appreciate the opportunity to provide these services to you. Please do not hesitate to contact me if you have any questions.'
    run = para.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Sincerely
    para = doc.add_paragraph()
    run = para.add_run('Sincerely,')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # KIMLEY-HORN AND ASSOCIATES, INC.
    para = doc.add_paragraph()
    run = para.add_run('KIMLEY-HORN AND ASSOCIATES, INC.')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.bold = True
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # KH Signer Name
    para = doc.add_paragraph()
    kh_name = invoice_info.get('kh_signer_name', '[Name]')
    run = para.add_run(kh_name)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # KH Signer Title
    para = doc.add_paragraph()
    kh_title = invoice_info.get('kh_signer_title', '[Title]')
    run = para.add_run(kh_title)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Client Name (legal entity)
    para = doc.add_paragraph()
    run = para.add_run(client_info.get('legal_entity', client_info.get('name', '')).upper())
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.bold = True
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Client signature lines
    para = doc.add_paragraph()
    run = para.add_run('SIGNED: _______________________________')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    para = doc.add_paragraph()
    run = para.add_run('PRINTED NAME: _______________________')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    para = doc.add_paragraph()
    run = para.add_run('TITLE:_________________________________')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    para = doc.add_paragraph()
    run = para.add_run('DATE: _______________________________')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.0
    
    doc.add_paragraph()
    
    # Attachment reference
    para = doc.add_paragraph()
    run = para.add_run('Attachment – Standard Provisions')
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.italic = True
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0


def generate_proposal_document(client_info, project_info, selected_tasks, assumptions, permits, invoice_info, included_additional_services, included_additional_services_with_fees, excluded_additional_services, output_path):
    """Generate complete proposal document."""
    
    doc = Document()
    
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    create_header(section)
    create_footer(section)
    
    add_opening_section(doc, client_info, project_info)
    add_project_understanding(doc, project_info, assumptions)
    add_scope_of_services(doc, selected_tasks, permits, included_additional_services)
    add_scope_table(doc, selected_tasks, included_additional_services_with_fees)
    add_information_provided_section(doc)
    add_client_responsibilities_section(doc)
    add_schedule_section(doc)
    add_fee_and_billing_section(doc, selected_tasks)
    add_additional_services_section(doc, excluded_additional_services)
    add_closure_section(doc, client_info, invoice_info)
    
    doc.save(output_path)
    return output_path

# ============================================================================
# STREAMLIT APP
# ============================================================================

st.set_page_config(
    page_title="Development Services Proposal Generator",
    page_icon="🏗️",
    layout="wide"
)

# ============================================================================
# STREAMLIT APP
# ============================================================================

st.set_page_config(
    page_title="Development Services Proposal Generator",
    page_icon="🏗️",
    layout="wide"
)

st.title("🏗️ Development Services Proposal Generator")
st.markdown("---")

# Create tabs
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📍 Property & Client",
    "📋 Project & Assumptions",
    "✅ Scope of Services",
    "📧 Permitting & Summary",
    "💰 Invoice & Billing"
])


# ============================================================================
# STREAMLIT UI WITH TABS
# ============================================================================

st.title("📋 Development Services Proposal Generator")
st.caption("Kimley-Horn - Tampa Bay Counties (Pinellas, Hillsborough, Pasco)")

tab1, tab2, tab3, tab4 = st.tabs([
    "🏠 Property & Client Info",
    "📋 Scope of Services",
    "✓ Assumptions & Permits",
    "💰 Invoice & Generate"
])

# TAB 1: Property Lookup & Client Info
with tab1:
    st.markdown("---")
    
    # County Selection
    county = st.selectbox(
        "Select County",
        ["Pinellas", "Hillsborough", "Pasco"],
        help="Choose the county where your property is located"
    )
    
    # Input Section
    st.subheader("Input")
    
    if county == "Pinellas":
        parcel_id_input = st.text_input(
            "Parcel ID",
            placeholder="e.g., 19-31-17-73166-001-0010",
            help="Pinellas County parcel ID with dashes",
            key="parcel_input"
        )
    elif county == "Hillsborough":
        parcel_id_input = st.text_input(
            "Folio Number",
            placeholder="e.g., 109054.1000",
            help="Hillsborough County folio number",
            key="parcel_input"
        )
    else:  # Pasco
        parcel_id_input = st.text_input(
            "Parcel ID",
            placeholder="e.g., 29-24-17-0000-0D411-0000",
            help="Pasco County parcel ID (with or without dashes)",
            key="parcel_input"
        )
    
    # Lookup Buttons
    if st.button("🔍 Lookup Property Info", type="primary"):
        if not parcel_id_input:
            st.error(f"Please enter a {'parcel ID' if county == 'Pinellas' else 'folio number'}")
        else:
            try:
                if county == "Pinellas":
                    # Validate
                    is_valid, error_msg = validate_parcel_id(parcel_id_input)
                    if not is_valid:
                        st.error(f"❌ {error_msg}")
                    else:
                        with st.spinner("Fetching property data from PCPAO API..."):
                            result = scrape_pinellas_property(parcel_id_input)
                            
                            if result['success']:
                                st.session_state['api_address'] = result.get('address', '')
                                st.session_state['api_city'] = result.get('city', '')
                                st.session_state['api_zip'] = result.get('zip', '')
                                st.session_state['api_owner'] = result.get('owner', '')
                                st.session_state['api_land_use'] = result.get('land_use', '')
                                st.session_state['api_zoning'] = result.get('zoning', '')
                                st.session_state['land_area_sqft'] = result.get('site_area_sqft', '')
                                st.session_state['land_area_acres'] = result.get('site_area_acres', '')
                                st.session_state['api_flu'] = ''
                                
                                st.success("✅ Property data retrieved successfully!")
                                st.rerun()
                            else:
                                st.error(f"❌ {result['error']}")
                
                elif county == "Hillsborough":
                    with st.spinner("Fetching property data from Hillsborough County..."):
                        result = lookup_hillsborough_parcel(parcel_id_input)
                        
                        if result['success']:
                            st.session_state['api_address'] = result.get('address', '')
                            st.session_state['api_city'] = result.get('city', '')
                            st.session_state['api_zip'] = result.get('zip', '')
                            st.session_state['api_owner'] = result.get('owner', '')
                            st.session_state['api_land_use'] = result.get('land_use', '')
                            st.session_state['api_zoning'] = result.get('zoning', '')
                            st.session_state['land_area_sqft'] = result.get('site_area_sqft', '')
                            st.session_state['land_area_acres'] = result.get('site_area_acres', '')
                            st.session_state['api_flu'] = ''
                            st.session_state['hillsborough_geometry'] = result.get('geometry')
                            
                            st.success("✅ Property data retrieved successfully!")
                            st.rerun()
                        else:
                            st.error(f"❌ {result['error']}")
                
                else:  # Pasco
                    
                    with st.spinner("Fetching property data from Pasco County (SWFWMD)..."):
                        result = lookup_pasco_parcel(parcel_id_input)
                        
                        
                        if result['success']:
                            st.session_state['api_address'] = result.get('address', '')
                            st.session_state['api_city'] = result.get('city', '')
                            st.session_state['api_zip'] = result.get('zip', '')
                            st.session_state['api_owner'] = result.get('owner', '')
                            st.session_state['api_land_use'] = result.get('land_use', '')
                            st.session_state['api_zoning'] = result.get('zoning', '')
                            st.session_state['land_area_sqft'] = result.get('site_area_sqft', '')
                            st.session_state['land_area_acres'] = result.get('site_area_acres', '')
                            st.session_state['api_flu'] = result.get('flu', '')  # DOR4CODE for Pasco
                            st.session_state['pasco_geometry'] = result.get('geometry')
                            
                            st.success("✅ Property data retrieved successfully!")
                            st.rerun()
                        else:
                            st.error(f"❌ {result.get('error', 'Unknown error')}")
            except Exception as e:
                st.error(f"❌ Unexpected error: {str(e)}")
                import traceback
                st.code(traceback.format_exc())
    
    # Second button for zoning/FLU lookup (always show after parcel ID entered)
    st.markdown("---")
    
    if county == "Pinellas":
        st.caption("🔍 **For Pinellas County:** Lookup detailed zoning and Future Land Use from GIS layers (St. Petersburg, Clearwater, Largo, and unincorporated areas)")
    elif county == "Hillsborough":
        st.caption("🔍 **For Hillsborough County:** Lookup detailed zoning and Future Land Use from county GIS layers")
    else:  # Pasco
        st.caption("🔍 **For Pasco County:** Lookup detailed zoning and Future Land Use from county GIS layers")
    
    if st.button("🗺️ Lookup Zoning & Future Land Use", type="secondary"):
        city = st.session_state.get('api_city', '')
        address = st.session_state.get('api_address', '')
        
        if not address:
            st.error("❌ Please run Property Lookup first to get the address")
        else:
            if county == "Pinellas":
                with st.spinner(f"Fetching zoning data for {address} in {city}..."):
                    zoning_result = lookup_pinellas_zoning(city, address)
            
            elif county == "Hillsborough":
                geom = st.session_state.get('hillsborough_geometry')
                with st.spinner(f"Fetching zoning/FLU data for {address}..."):
                    zoning_result = lookup_hillsborough_zoning_flu(address, geometry=geom)
            
            else:  # Pasco
                geom = st.session_state.get('pasco_geometry')
                with st.spinner(f"Fetching zoning/FLU data for {address}..."):
                    zoning_result = lookup_pasco_zoning_flu(address, geometry=geom)
            
            if zoning_result.get('success'):
                # Update zoning with detailed info
                if zoning_result.get('zoning_code'):
                    if zoning_result.get('zoning_description'):
                        st.session_state['api_zoning'] = f"{zoning_result.get('zoning_code')} - {zoning_result.get('zoning_description')}"
                    else:
                        st.session_state['api_zoning'] = zoning_result.get('zoning_code', '')
                
                # Update FLU
                if zoning_result.get('future_land_use'):
                    if zoning_result.get('future_land_use_description'):
                        st.session_state['api_flu'] = f"{zoning_result.get('future_land_use')} - {zoning_result.get('future_land_use_description')}"
                    else:
                        st.session_state['api_flu'] = zoning_result.get('future_land_use', '')
                
                # Show which jurisdiction was queried
                if county == "Pasco":
                    st.success(f"✅ Zoning and FLU data updated from Pasco County GIS!")
                elif county == "Hillsborough":
                    st.success(f"✅ Zoning and FLU data updated from Hillsborough County GIS!")
                elif 'St. Petersburg' in city or 'St Petersburg' in city:
                    st.success(f"✅ Zoning data updated from St. Petersburg GIS layers!")
                elif 'Clearwater' in city:
                    st.success(f"✅ Zoning data updated from Clearwater GIS layers!")
                elif 'Largo' in city:
                    st.success(f"✅ Zoning/FLU data updated (Largo uses Future Land Use classification)!")
                elif 'Unincorporated' in city:
                    st.success(f"✅ Zoning data updated from Pinellas County (unincorporated) layers!")
                else:
                    st.success(f"✅ Zoning data updated!")
                
                st.rerun()
            else:
                st.error(f"❌ {zoning_result.get('error', 'Unable to fetch zoning data')}")
    
    st.markdown("---")
    
    # Results Section
    st.subheader("Results - All Property Data")
    st.caption("These fields auto-fill after successful lookup")
    
    col_left, col_right = st.columns(2)
    
    with col_left:
        st.text_input(
            "City (auto-filled)",
            key='api_city',
            placeholder="Will auto-fill",
            help="City name from Property Appraiser"
        )
        
        st.text_input(
            "Address (auto-filled)",
            key='api_address',
            placeholder="Will auto-fill",
            help="Property address from Property Appraiser"
        )
        
        st.text_input(
            "ZIP Code (auto-filled)",
            key='api_zip',
            placeholder="Will auto-fill",
            help="ZIP code from Property Appraiser"
        )
        
        st.text_input(
            "Property Use (auto-filled)",
            key='api_land_use',
            placeholder="Will auto-fill",
            help="Property Appraiser land use classification"
        )
    
    with col_right:
        st.text_input(
            "Owner (auto-filled)",
            key='api_owner',
            placeholder="Will auto-fill",
            help="Property owner from Property Appraiser"
        )
        
        st.text_input(
            "Zoning (auto-filled)",
            key='api_zoning',
            placeholder="Will auto-fill from GIS layers",
            help="Zoning district from county GIS layers"
        )
        
        st.text_input(
            "Future Land Use (auto-filled)",
            key='api_flu',
            placeholder="Will auto-fill from GIS layers",
            help="Future Land Use from county GIS layers"
        )
        
        st.text_input(
            "Land Area (acres)",
            key='land_area_acres',
            placeholder="Will auto-fill",
            help="Acreage from Property Appraiser"
        )
    
    st.text_input(
        "Land Area (square feet)",
        key='land_area_sqft',
        placeholder="Will auto-fill (Pinellas only)",
        help="Square footage from Property Appraiser (Pinellas only)"
    )
    
    # Summary
    st.markdown("---")
    st.subheader("Test Summary")
    
    if st.session_state.get('api_city'):
        st.success("✅ Step 1: PCPAO API Lookup completed")
        
        # Show what was retrieved
        retrieved = []
        if st.session_state.get('api_city'): retrieved.append("City")
        if st.session_state.get('api_address'): retrieved.append("Address")
        if st.session_state.get('api_owner'): retrieved.append("Owner")
        if st.session_state.get('api_land_use'): retrieved.append("Property Use")
        if st.session_state.get('land_area_acres'): retrieved.append("Land Area")
        
        st.info(f"**Retrieved from PCPAO:** {', '.join(retrieved)}")
        
        # Check if zoning/FLU have been looked up
        if st.session_state.get('api_flu'):
            st.success("✅ Step 2: GIS Layer Lookup completed")
            zoning_retrieved = []
            if st.session_state.get('api_zoning') and 'Contact City' not in st.session_state.get('api_zoning', ''): 
                zoning_retrieved.append("Zoning")
            if st.session_state.get('api_flu'): 
                zoning_retrieved.append("Future Land Use")
            if zoning_retrieved:
                st.info(f"**Retrieved from GIS Layers:** {', '.join(zoning_retrieved)}")
        else:
            st.info("ℹ️ Click '🗺️ Lookup Zoning & Future Land Use' button to get detailed zoning data (St. Petersburg only)")
    else:
        st.info("Click '🔍 Lookup Property Info' to start")
# TAB 2: Scope of Services
with tab2:
        st.subheader("Scope of Services")
        st.markdown("Select the tasks to include, enter the fee, and choose the fee type for each task.")
        
        selected_tasks = {}
        
        for task_num in sorted(DEFAULT_FEES.keys()):
            task = DEFAULT_FEES[task_num]
            
            col_check, col_name, col_fee, col_type = st.columns([0.5, 3, 1.5, 1.5])
            
            with col_check:
                task_selected = st.checkbox(
                    f"{task_num}",
                    value=(task_num == '310'),  # Auto-check Task 310
                    key=f"check_{task_num}",
                    label_visibility="collapsed"
                )
            
            with col_name:
                if task_num == '310':
                    st.markdown(f"**Task {task_num}: {task['name']}** *(uncheck if not needed)*")
                else:
                    st.markdown(f"**Task {task_num}: {task['name']}**")
            
            with col_fee:
                fee_amount = st.number_input(
                    "Fee ($)",
                    min_value=0,
                    value=None,
                    placeholder=f"{task['amount']:,}",
                    key=f"fee_{task_num}",
                    disabled=not task_selected,
                    label_visibility="collapsed"
                )
            
            with col_type:
                fee_type_selection = st.selectbox(
                    "Type",
                    options=["Hourly, Not-to-Exceed", "Hourly", "Lump Sum"],
                    index=0,  # Default to Hourly, Not-to-Exceed
                    key=f"type_{task_num}",
                    disabled=not task_selected,
                    label_visibility="collapsed"
                )
            
            if task_selected:
                final_fee = fee_amount if fee_amount is not None else task['amount']
                selected_tasks[task_num] = {
                    'name': task['name'],
                    'fee': final_fee,
                    'type': fee_type_selection  # Use per-task fee type
                }
                
                # Task 310 - Construction Phase Services selection
                if task_num == '310':
                    st.markdown("**📋 Construction Phase Services:**")
                    st.caption("Select services, enter hours/count, rate, and cost")
                    
                    # Header row
                    col_h1, col_h2, col_h3, col_h4, col_h5 = st.columns([0.5, 3, 1.5, 1.5, 1.5])
                    with col_h1:
                        st.write("")
                    with col_h2:
                        st.markdown("**Service**")
                    with col_h3:
                        st.markdown("**Hrs/Count**")
                    with col_h4:
                        st.markdown("**$/hr**")
                    with col_h5:
                        st.markdown("**Cost ($)**")
                    
                    # Services configuration
                    services_list = [
                        ('shop_drawings', 'Shop Drawing Review', 30, 165, 4950),
                        ('rfi', 'RFI Response', 50, 165, 8250),
                        ('oac', 'OAC Meetings', 24, 0, 3000),
                        ('site_visits', 'Site Visits (2 hrs each)', 4, 0, 1000),
                        ('asbuilt', 'As-Built Reviews', 2, 0, 500),
                        ('inspection_tv', 'Inspection & TV Reports', 0, 165, 0),
                        ('record_drawings', 'Record Drawings (Water/Sewer)', 40, 165, 6600),
                        ('fdep', 'FDEP Clearance Submittals', 0, 0, 0),
                        ('compliance', 'Letter of General Compliance', 0, 0, 0),
                        ('wmd', 'WMD Certification', 0, 0, 0)
                    ]
                    
                    service_data = {}
                    
                    for svc_key, svc_name, default_hrs, default_rate, default_cost in services_list:
                        col_chk, col_nm, col_hrs, col_rate, col_cost = st.columns([0.5, 3, 1.5, 1.5, 1.5])
                        
                        with col_chk:
                            is_selected = st.checkbox(
                                "✓",
                                value=(svc_key in ['shop_drawings', 'rfi', 'oac', 'site_visits', 'asbuilt', 'fdep', 'compliance', 'wmd']),
                                key=f"svc310_{svc_key}",
                                label_visibility="collapsed"
                            )
                        
                        with col_nm:
                            st.markdown(f"{svc_name}")
                        
                        with col_hrs:
                            if default_hrs > 0 or svc_key in ['inspection_tv', 'record_drawings']:
                                hrs_value = st.number_input(
                                    "Hrs",
                                    min_value=0,
                                    value=default_hrs,
                                    key=f"hrs310_{svc_key}",
                                    disabled=not is_selected,
                                    label_visibility="collapsed"
                                )
                            else:
                                hrs_value = 0
                                st.write("—")
                        
                        with col_rate:
                            if default_rate > 0 or svc_key in ['inspection_tv', 'record_drawings']:
                                rate_value = st.number_input(
                                    "Rate",
                                    min_value=0,
                                    value=default_rate,
                                    key=f"rate310_{svc_key}",
                                    disabled=not is_selected,
                                    label_visibility="collapsed"
                                )
                            else:
                                rate_value = 0
                                st.write("—")
                        
                        with col_cost:
                            if is_selected:
                                cost_value = st.number_input(
                                    "Cost",
                                    min_value=0,
                                    value=default_cost,
                                    key=f"cost310_{svc_key}",
                                    disabled=not is_selected,
                                    label_visibility="collapsed"
                                )
                            else:
                                cost_value = 0
                                st.write("—")
                        
                        service_data[svc_key] = {
                            'included': is_selected,
                            'name': svc_name,
                            'hours': hrs_value if is_selected else 0,
                            'rate': rate_value if is_selected else 0,
                            'cost': cost_value if is_selected else 0
                        }
                    
                    st.markdown("---")
                    total_hrs = st.number_input(
                        "**Total Task 310 Hours**",
                        min_value=0,
                        value=180,
                        key="total_construction_hours"
                    )
                    
                    selected_tasks[task_num]['services'] = service_data
                    selected_tasks[task_num]['total_hours'] = total_hrs
                    
                    # Create hours dictionary for placeholder replacement in document generation
                    selected_tasks[task_num]['hours'] = {
                        'shop_drawing': service_data['shop_drawings']['hours'],
                        'rfi': service_data['rfi']['hours'],
                        'oac_meetings': service_data['oac']['hours'],
                        'site_visits': service_data['site_visits']['hours'],
                        'record_drawing': service_data['record_drawings']['hours'],
                        'total': total_hrs
                    }
    
    
    
    

# TAB 3: Assumptions & Permits
with tab3:
        st.subheader("Permitting Requirements")
        st.markdown("Select the permits/approvals required for this project (applies to Task 150 - Civil Permitting):")
        
        permit_config = PERMIT_MAPPING.get(st.session_state.get('county', ''), {})
        default_permits = permit_config.get('default_permits', [])
        ahj_name = permit_config.get('ahj_name', 'Authority Having Jurisdiction')
        wmd_name = permit_config.get('wmd_short', 'Water Management District')
        
        col_permit1, col_permit2, col_permit3 = st.columns(3)
        
        with col_permit1:
            permit_ahj = st.checkbox(f"{ahj_name}", value=("ahj" in default_permits), help="Primary permitting authority")
            permit_sewer = st.checkbox("Sewer Provider", value=("sewer" in default_permits))
            permit_water = st.checkbox("Water Provider", value=("water" in default_permits))
        
        with col_permit2:
            permit_wmd_erp = st.checkbox(f"{wmd_name} ERP", value=("wmd_erp" in default_permits), help="Environmental Resources Permit")
            permit_fdep = st.checkbox("FDEP Potable Water/Wastewater", value=("fdep" in default_permits))
            permit_fdot_drainage = st.checkbox("FDOT Drainage Connection", value=("fdot_drainage" in default_permits))
        
        with col_permit3:
            permit_fdot_driveway = st.checkbox("FDOT Driveway Connection", value=("fdot_driveway" in default_permits))
            permit_fdot_utility = st.checkbox("FDOT Utility Connection", value=("fdot_utility" in default_permits))
            permit_fema = st.checkbox("FEMA", value=("fema" in default_permits))
        
        st.markdown("---")
        
        # Additional Services Configuration
        st.subheader("Additional Services")
        st.markdown("**Check the services you ARE providing** in this proposal and enter the fee. Unchecked services will be listed as 'Additional Services' (not included).")
        
        # Default list of additional services with suggested default fees
        additional_services_list = [
            ("offsite_roadway", "Off-site roadway, traffic signal design or utility improvements", False, 25000),
            ("offsite_utility", "Off-site utility capacity analysis and extensions", False, 15000),
            ("utility_relocation", "Utility relocation design and plans", False, 12000),
            ("cost_opinions", "Preparation of opinions of probable construction costs", False, 5000),
            ("dewatering", "Dewatering permitting (to be provided by Contractor)", False, 3000),
            ("site_lighting", "Site lighting, photometric, and site electrical plan", False, 8000),
            ("dry_utility", "Dry utility coordination and design", False, 10000),
            ("landscape", "Landscape, irrigation, hardscape design and tree mitigation", False, 20000),
            ("fire_line", "Fire line design", False, 6000),
            ("row_permitting", "Right-of-way permitting", False, 8000),
            ("concurrency", "Concurrency application assistance", False, 5000),
            ("3d_modeling", "3D modeling and graphic/presentations", False, 8000),
            ("leed", "LEED certification and review", False, 20000),
            ("schematic_dd", "Schematic and design development plans", False, 15000),
            ("extra_meetings", "Meetings other than those described in the tasks above", False, 5000),
            ("surveying", "Boundary, topographic and tree surveying, platting and subsurface utility exploration", False, 25000),
            ("platting", "Platting or easement assistance", False, 8000),
            ("traffic_studies", "Traffic studies, analysis, property share agreement", False, 30000),
            ("mot_plans", "Maintenance of traffic plans", False, 12000),
            ("structural", "Structural engineering (including retaining walls)", False, 35000),
            ("signage", "Signage design", False, 4000),
            ("extra_design", "Design elements beyond those outlined in the above project understanding", False, 10000),
            ("peer_review", "Responding to comments from third-party peer review", False, 8000)
        ]
        
        st.caption("💡 Tip: Check services you ARE including and enter fees. Unchecked items appear in 'Additional Services (Not Included)' section.")
        
        # Create 2-column layout: Service name (left) + Fee (right)
        included_additional_services = []
        excluded_additional_services = []
        included_additional_services_with_fees = {}  # Store as dict with fees
        
        for key, service_name, default_checked, default_fee in additional_services_list:
            col_service, col_fee = st.columns([3, 1])
            
            with col_service:
                is_checked = st.checkbox(
                    service_name,
                    value=default_checked,
                    key=f"addl_svc_{key}"
                )
            
            with col_fee:
                fee_amount = st.number_input(
                    "Fee ($)",
                    min_value=0,
                    value=None,
                    placeholder=f"{default_fee:,}",
                    key=f"addl_fee_{key}",
                    disabled=not is_checked,
                    label_visibility="collapsed"
                )
            
            if is_checked:
                final_fee = fee_amount if fee_amount is not None else default_fee
                included_additional_services.append(service_name)
                included_additional_services_with_fees[service_name] = final_fee
            else:
                excluded_additional_services.append(service_name)
        
        st.markdown("---")
        
        # Summary
        st.subheader("Selected Tasks Summary")
        if selected_tasks or included_additional_services_with_fees:
            total_fee = 0
            
            # Show regular tasks
            for task_num in sorted(selected_tasks.keys()):
                task = selected_tasks[task_num]
                st.write(f"✓ Task {task_num}: {task['name']} — **${task['fee']:,}**")
                total_fee += task['fee']
            
            # Show included additional services
            if included_additional_services_with_fees:
                st.markdown("**Additional Services Included:**")
                for service_name, service_fee in included_additional_services_with_fees.items():
                    st.write(f"✓ {service_name} — **${service_fee:,}**")
                    total_fee += service_fee
            
            st.markdown("---")
            st.markdown(f"### **Total Fee: ${total_fee:,}**")
        else:
            st.info("👆 Select at least one task in the Scope of Services tab")
    
    

# TAB 4: Invoice & Generate
with tab4:
        st.subheader("Invoice & Billing Information")
        col_inv1, col_inv2 = st.columns(2)
        
        with col_inv1:
            invoice_email = st.text_input(
                "Invoice Email Address",
                placeholder="e.g., accounting@company.com",
                help="Primary email for invoices"
            )
            kh_signer_name = st.text_input(
                "Kimley-Horn Signer Name",
                placeholder="e.g., John Smith, PE"
            )
            
            # Retainer checkbox and amount
            use_retainer = st.checkbox(
                "Require Retainer",
                value=False,
                help="Check if this proposal requires an upfront retainer fee"
            )
        
        with col_inv2:
            invoice_cc_email = st.text_input(
                "CC Email (optional)",
                placeholder="e.g., manager@company.com",
                help="Additional recipient for invoices"
            )
            kh_signer_title = st.text_input(
                "Kimley-Horn Signer Title",
                placeholder="e.g., Senior Project Manager"
            )
            
            # Retainer amount (disabled if not using retainer)
            retainer_amount = st.number_input(
                "Retainer Amount ($)",
                min_value=0,
                value=0,
                disabled=not use_retainer,
                help="Upfront retainer fee required before work begins"
            )
        
        st.markdown("---")
        st.markdown("---")
        
        # Generate Document Section
        st.subheader("📄 Generate Proposal Document")
        
        required_fields = {
            'County': st.session_state.get('county', ''),
            'City': st.session_state.get('city', ''),
            'Client Name': st.session_state.get('client_name', ''),
            'Legal Entity Name': st.session_state.get('legal_entity_name', ''),
            'Contact Person': st.session_state.get('contact_person', ''),
            'Address Line 1': st.session_state.get('address_line1', ''),
            'Address Line 2': st.session_state.get('address_line2', ''),
            'Project Name': st.session_state.get('project_name', ''),
            'Project Description': st.session_state.get('project_description', '')
        }
        
        missing_fields = [field for field, value in required_fields.items() if not value]
        
        if missing_fields:
            st.warning(f"⚠️ Please fill in: {', '.join(missing_fields)}")
        
        if not selected_tasks:
            st.warning("⚠️ Please select at least one task in the Scope of Services tab")
        
        can_generate = not missing_fields and bool(selected_tasks)
        
        if st.button("🚀 Generate Proposal Document", type="primary", disabled=not can_generate):
            with st.spinner("Generating proposal document..."):
                try:
                    # Collect assumptions
                    assumptions = []
                    if assume_survey:
                        assumptions.append("Boundary, topographic, and tree survey will be provided by the Client.")
                    if assume_environmental:
                        assumptions.append("An Environmental/Biological assessment and Geotechnical investigation report will be provided by the Client.")
                    if assume_geotech:
                        assumptions.append("A Geotechnical investigation report will be provided by the Client.")
                    if assume_zoning:
                        assumptions.append("The proposed use is consistent with the property's future land use and zoning designations.")
                    if has_conceptual_plan and conceptual_plan_date:
                        assumptions.append(f"This proposal is based on the conceptual site plan dated {conceptual_plan_date}.")
                    if assume_utilities:
                        assumptions.append("Utilities are available at the project boundary and have the capacity to serve the proposed development.")
                    if assume_offsite:
                        assumptions.append("Offsite roadway improvements or right-of-way permitting is not included.")
                    if assume_traffic:
                        assumptions.append("Traffic Study, impact analysis, and traffic counts, if required, will be provided by others.")
                    if assume_one_phase:
                        assumptions.append("The project will be constructed in one (1) phase.")
                    
                    # Collect permitting requirements
                    permit_config = PERMIT_MAPPING.get(st.session_state.get('county', ''), {})
                    ahj_name = permit_config.get('ahj_name', 'Authority Having Jurisdiction')
                    wmd_full = permit_config.get('wmd', 'Water Management District')
                    
                    permits = []
                    if permit_ahj:
                        permits.append(ahj_name)
                    if permit_sewer:
                        permits.append(f"{ahj_name} Sewer")
                    if permit_water:
                        permits.append(f"{ahj_name} Water")
                    if permit_wmd_erp:
                        permits.append(f"{wmd_full} Environmental Resources Permit (ERP)")
                    if permit_fdep:
                        permits.append("Florida Department of Environmental Protection (FDEP) Potable Water and Wastewater Permit")
                    if permit_fdot_drainage:
                        permits.append("FDOT Drainage Connection Permit")
                    if permit_fdot_driveway:
                        permits.append("FDOT Driveway Connection Permit")
                    if permit_fdot_utility:
                        permits.append("FDOT Utility Connection Permit")
                    if permit_fema:
                        permits.append("FEMA")
                    
                    client_info = {
                        'name': st.session_state.get('client_name', ''),
                        'legal_entity': st.session_state.get('legal_entity_name', ''),
                        'contact': st.session_state.get('contact_person', ''),
                        'address1': st.session_state.get('address_line1', ''),
                        'address2': st.session_state.get('address_line2', ''),
                        'phone': st.session_state.get('phone', ''),
                        'email': st.session_state.get('email', '')
                    }
                    
                    project_info = {
                        'date': st.session_state.get('proposal_date', date.today()).strftime("%B %d, %Y"),
                        'name': st.session_state.get('project_name', ''),
                        'address': st.session_state.get('project_address', ''),
                        'city_state_zip': st.session_state.get('project_city_state_zip', ''),
                        'description': st.session_state.get('project_description', ''),
                        'county': st.session_state.get('county', ''),
                        'city': st.session_state.get('city', ''),
                        'parcel_id': st.session_state.get('parcel_id', ''),
                        'permits': permits,
                        # Property lookup data for Project Understanding paragraph
                        'site_acres': st.session_state.get('lookup_site_area_acres', ''),
                        'future_land_use': st.session_state.get('lookup_land_use', ''),
                        'zoning': st.session_state.get('lookup_zoning', '')
                    }
                    
                    invoice_info = {
                        'email': invoice_email,
                        'cc_email': invoice_cc_email,
                        'kh_signer_name': kh_signer_name,
                        'kh_signer_title': kh_signer_title,
                        'use_retainer': use_retainer,
                        'retainer_amount': retainer_amount if use_retainer else 0
                    }
                    
                    buffer = BytesIO()
                    temp_path = '/tmp/temp_proposal.docx'
                    generate_proposal_document(client_info, project_info, selected_tasks, assumptions, permits, invoice_info, included_additional_services, included_additional_services_with_fees, excluded_additional_services, temp_path)
                    
                    with open(temp_path, 'rb') as f:
                        buffer.write(f.read())
                    buffer.seek(0)
                    
                    filename = f"Proposal_{st.session_state.get('project_name', 'Document').replace(' ', '_')[:30]}_{st.session_state.get('proposal_date', date.today()).strftime('%Y%m%d')}.docx"
                    
                    st.success("✅ **Proposal document generated successfully!**")
                    
                    st.download_button(
                        label="📥 Download Word Document",
                        data=buffer.getvalue(),
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary",
                        use_container_width=True
                    )
                    
                except Exception as e:
                    st.error(f"❌ **Error:** {str(e)}")
                    with st.expander("Show Error Details"):
                        st.exception(e)

st.markdown("---")
st.caption("Development Services Proposal Generator | Kimley-Horn - Tampa Bay Counties")
